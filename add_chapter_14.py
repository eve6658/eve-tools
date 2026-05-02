#!/usr/bin/env python3
"""给交易手册添加第十四章：小单占比经验阈值与实盘标注"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime

input_file = '/home/adam/.openclaw/workspace/A股不对称战场操盘手册_完整版_含第十三章.docx'
output_file = '/home/adam/.openclaw/workspace/A股不对称战场操盘手册_完整版_含第十四章.docx'

doc = Document(input_file)

def add_heading_styled(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(139, 0, 0) if level == 1 else RGBColor(0, 51, 102)
    return p

def add_para(text='', bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_table(data):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
    return t

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ===== 插入分页 + 新章节 =====
doc.add_page_break()

add_heading_styled('第十四章  小单占比经验阈值 — 盘口微观信号量化研究', level=1)

add_para('本章基于2026年4月2日两笔实操的Level-2千档盘口逐笔成交数据，'
         '将"小单占比判断主力行为"这一知识点从定性认知推进到定量研究阶段。')

add_para('数据来源：财达股市通 Level-2 千档盘口 + 逐笔成交明细',
         bold=True, color=RGBColor(100, 100, 100))

# 14.1
add_heading_styled('14.1 研究背景与核心问题', level=2)

add_para('在前面的章节中，我们已经建立了"大单流向判断主力行为"的框架。'
         '但有一个关键盲区：主力可以拆单伪装——把一笔500手的大单拆成10笔50手的小单，'
         '让资金流向数据显示为"小单净流入"，从而隐藏真实意图。')

add_para('这就引出一个核心问题：', bold=True)
add_para('• 如何区分"真正的小单（散户）"和"伪装的小单（主力拆单）"？')
add_para('• 小单占比在什么阈值以上，可以判断为"主力在用小单操作"？')
add_para('• 这个阈值在建仓、出货、洗盘等不同场景下是否不同？')

add_para()
add_para('核心命题：', bold=True, color=RGBColor(139, 0, 0))
add_para('"小单占比异常升高是主力拆单操作的必要不充分条件。'
         '必须结合价格区间稳定性、量能方向、压单/托单行为等维度综合判断。"',
         bold=True)

# 14.2
add_heading_styled('14.2 关键概念定义', level=2)

add_para('在开始分析之前，先统一口径：')

concept_data = [
    ['概念', '定义', '说明'],
    ['小单', '单笔成交 < 100手', '当前假设阈值，需后续验证'],
    ['中单', '单笔成交 100-500手', '可能是中等资金或主力碎片单'],
    ['大单', '单笔成交 > 500手', '主力特征明显'],
    ['小单占比', '小单笔数 / 总成交笔数 × 100%', '核心观测指标'],
    ['散户典型手数', '20-50手/笔', '基于实盘逐笔统计'],
    ['主力典型手数', '100-500手/笔', '基于实盘逐笔统计'],
    ['伪装小单', '主力拆单后的小单', '特征：异常均匀、持续不断'],
]
add_table(concept_data)

# 14.3
add_heading_styled('14.3 案例一：600666 奥瑞德 — 压盘吸筹型', level=2)

add_para('2026年4月2日，奥瑞德全天走势呈现典型的"压盘吸筹"特征：'
         '开盘6.27，快速拉升至6.65后缩量回落，全天在6.10-6.30窄幅震荡。')

add_heading_styled('14.3.1 逐笔成交小单分析', level=3)

add_para('关键时段：10:17-10:24（回落企稳阶段）', bold=True)

order_data = [
    ['档位', '平均手数/笔', '特征', '判断'],
    ['卖一档（6.18-6.19）', '39手', '碎片化、连续不断', '散户卖单'],
    ['买一档（6.10-6.12）', '152手', '间歇性大单承接', '主力买单'],
]
add_table(order_data)

add_para()
add_para('解读：', bold=True)
add_para('卖一档平均39手/笔，属于典型的散户手数区间（20-50手）。'
         '买一档平均152手/笔，属于主力手数区间（100-500手）。'
         '散户在卖，主力在买——这就是"压盘吸筹"的微观证据。')

add_heading_styled('14.3.2 盘口铁底全天未动', level=3)

add_para('千档盘口底部筹码全天纹丝不动，这是主力控盘的铁证：')

bottom_data = [
    ['价位', '挂单量（手）', '变化'],
    ['6.07', '5,031', '全天未动'],
    ['6.05', '3,924', '全天未动'],
    ['6.03', '4,020', '全天未动'],
    ['6.02', '4,013', '全天未动'],
    ['6.00', '3,133', '全天未动'],
    ['合计', '25,000+', '铁底'],
]
add_table(bottom_data)

add_heading_styled('14.3.3 卖盘变化时间线', level=3)

timeline_data = [
    ['时间', '卖盘总量', '买盘总量', '变化', '判断'],
    ['10:15', '48万手', '23万手', '卖>买', '主力压盘'],
    ['11:25', '37万手', '37万手', '均衡化', '主力撤压单'],
    ['14:17', '卖1: 150手', '买1: 2175手', '压单撤退', '即将拉升'],
    ['14:47', '477手（全部卖盘）', '25,000+手', '压单撤光', '确认信号'],
]
add_table(timeline_data)

add_heading_styled('14.3.4 小单占比量化结果', level=3)

result_data = [
    ['指标', '数值', '含义'],
    ['卖盘小单占比', '> 80% 🔥', '散户碎片化卖出'],
    ['买盘小单占比', '< 30%', '主力大单承接'],
    ['价格区间', '6.10-6.18（振幅1.3%）', '窄幅震荡，价格稳定'],
    ['量能方向', '缩量回落', '不是出货形态'],
    ['铁底状态', '25,000+手未动', '主力控盘'],
]
add_table(result_data)

# 14.4
add_heading_styled('14.4 案例二：002788 鹭燕医药 — 尾盘抢筹型', level=2)

add_para('同日尾盘14:52-14:54，鹭燕医药出现连续大单买入，呈现"尾盘抢筹"特征。')

add_heading_styled('14.4.1 逐笔成交明细', level=3)

tick_data = [
    ['时间', '方向', '价格', '手数', '判断'],
    ['14:54:03', 'B（买）', '16.97', '276手 🔥', '主力建仓'],
    ['14:53:58', 'B（买）', '16.97', '26手', '散户跟进'],
    ['14:53:37', 'B（买）', '16.98', '125手', '主力买入'],
    ['14:52:55', 'S（卖）', '16.98', '103手', '散户卖'],
    ['14:52:27', 'B（买）', '16.97', '265手', '主力买入'],
]
add_table(tick_data)

add_heading_styled('14.4.2 小单占比量化结果', level=3)

result2_data = [
    ['指标', '数值', '含义'],
    ['买盘平均手数', '165手', '主力特征'],
    ['小单占比（买盘尾盘）', '< 30% 🔥', '主力抢筹，非散户跟风'],
    ['价格区间', '16.97-16.98', '极窄幅，主力控价'],
    ['量能特征', '连续B大单', '确认主力进场'],
]
add_table(result2_data)

add_heading_styled('14.4.3 盘口压力与支撑', level=3)

orderbook_data = [
    ['档位', '卖价', '卖量（手）', '买价', '买量（手）'],
    ['5', '17.03', '78', '16.93', '223'],
    ['4', '17.02', '37', '16.92', '149'],
    ['3', '17.01', '130', '16.91', '68'],
    ['2', '17.00', '553 ⚠️', '16.90', '450'],
    ['1', '16.99', '146', '16.89', '69'],
]
add_table(orderbook_data)

add_para()
add_para('关键压力：17.00处553手压单 — 次日需要放量突破')
add_para('关键支撑：16.90处450手托单 — 如果跌破需警惕')

# 14.5
add_heading_styled('14.5 经验阈值初步结论', level=2)

add_para('基于以上两个案例的实盘数据，提出以下经验阈值假设：')

add_heading_styled('假设一：小单占比阈值', level=3)

threshold_data = [
    ['场景', '卖盘小单占比', '买盘小单占比', '信号'],
    ['吸筹（主力建仓）', '> 80% 🔥', '< 30%', '散户卖 + 主力买 = 吸筹'],
    ['出货（主力派发）', '< 30%', '> 80% 🔥', '主力卖 + 散户买 = 出货'],
    ['散户主导行情', '> 70%', '> 70%', '买卖都是散户 = 无方向'],
    ['主力对倒', '< 40%', '< 40%', '买卖都是主力 = 操纵'],
]
add_table(threshold_data)

add_heading_styled('假设二：必要不充分条件组合', level=3)

add_para('小单占比异常 + 以下任一条件 = 高置信度信号：')

add_bullet('✅ 价格区间稳定（振幅 < 2%）')
add_bullet('✅ 压单/托单反向运动（压单撤退 + 托单稳定）')
add_bullet('✅ 成交量方向一致（量增价稳 或 缩量回落）')
add_bullet('✅ 千档铁底未动（底部筹码纹丝不动）')

add_heading_styled('假设三：手数阈值', level=3)

hand_data = [
    ['分类', '手数区间', '典型特征', '判断'],
    ['小单', '< 100手', '散户碎片化', '可能散户，可能主力拆单'],
    ['中单', '100-500手', '间歇性出现', '主力或中等资金'],
    ['大单', '> 500手', '突爆性出现', '明确主力特征'],
    ['散户典型', '20-50手', '连续、均匀', '确认散户'],
    ['主力典型', '100-500手', '间歇、不均匀', '确认主力'],
]
add_table(hand_data)

# 14.6
add_heading_styled('14.6 实战操作流程', level=2)

add_para('将小单占比分析嵌入日常盯盘流程：', bold=True)

add_para('第一步：观察卖一档/买一档的平均手数', bold=True)
add_bullet('卖一档平均 < 50手 → 散户在卖')
add_bullet('买一档平均 > 100手 → 主力在买')
add_bullet('如果同时满足 → 吸筹信号')

add_para('第二步：验证价格区间稳定性', bold=True)
add_bullet('振幅 < 2% 且价格在窄幅箱体 → 价格稳定')
add_bullet('如果价格在急速下跌 → 小单占比高可能是恐慌抛售，不是吸筹')

add_para('第三步：检查压单/托单行为', bold=True)
add_bullet('卖盘压单频繁撤退 → 主力不想真卖，只想吓人')
add_bullet('买盘铁底纹丝不动 → 主力在托底')

add_para('第四步：综合判断', bold=True)
add_bullet('三个维度都确认 → 高置信度，可以操作')
add_bullet('只有小单占比异常，其他维度不确认 → 观望')

# 14.7
add_heading_styled('14.7 研究局限与后续方向', level=2)

add_para('当前研究的局限性：', bold=True)
add_bullet('样本量小：仅2个案例，统计意义有限')
add_bullet('数据源单一：仅财达Level-2，未交叉验证')
add_bullet('时间跨度短：仅单日数据，未覆盖不同市场环境')
add_bullet('阈值待验证：80%/30%是经验值，需更多样本确认')

add_para()
add_para('后续研究方向：', bold=True)
add_bullet('扩展案例：补充688347华虹公司、002378章源钨业、603938三孚股份等历史数据')
add_bullet('量化统计：统计各场景下小单占比分布，用统计方法验证阈值')
add_bullet('回测验证：用历史Level-2数据回测小单占比信号的有效性')
add_bullet('结合资金流向：将小单占比与东方财富大单/中单/小单数据交叉验证')
add_bullet('机器学习建模：用多维度特征训练主力行为分类模型')

# 14.8
add_heading_styled('14.8 本章要点总结', level=2)

summary_data = [
    ['要点', '内容'],
    ['核心发现', '卖盘小单占比>80%+买盘小单占比<30% → 吸筹信号'],
    ['必要条件', '小单占比异常是必要条件，不是充分条件'],
    ['验证维度', '价格稳定 + 压单撤退 + 铁底未动 + 量能方向'],
    ['实战流程', '观察手数→验证价格→检查盘口→综合判断'],
    ['数据来源', '2026-04-02 财达Level-2 千档盘口'],
    ['案例标的', '600666奥瑞德（吸筹）、002788鹭燕医药（抢筹）'],
]
add_table(summary_data)

add_para()
add_para('数据标注日期：2026年4月2日', color=RGBColor(120, 120, 120))
add_para('⚠️ 本章仅用于交易策略学习研究，不构成投资建议', color=RGBColor(180, 0, 0))

# 更新目录
doc.add_page_break()
add_para('（第十四章内容已插入手册末尾，建议重新生成目录以更新页码）',
         color=RGBColor(150, 150, 150))

# 保存
doc.save(output_file)
print(f'✅ 已添加第十四章：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
