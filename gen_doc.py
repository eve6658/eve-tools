from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 辅助函数：防止表格跨页
def make_table_keep_together(table):
    """表格及所有行设置不跨页"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.trPr if tr.trPr is not None else tr._add_trPr()
        for tag in ['w:keepNext', 'w:keepLines']:
            el = OxmlElement(tag)
            el.set(qn('w:val'), '1')
            trPr.append(el)

def add_styled_table(doc, headers, data, style_name='Light Grid Accent 1'):
    """添加一个不跨页的表格"""
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = style_name
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = str(val)
    make_table_keep_together(table)
    return table

# ============ 标题 ============
title = doc.add_heading('002788 鹭燕医药 & 600666 奥瑞德\n建仓计划书', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('生成时间：2026年4月2日 13:15 | 总资金：12,200元 | 策略：趋势跟踪+分歧建仓', style='Intense Quote')

# ============ 第一部分 ============
doc.add_heading('一、今日市场概况（2026.04.02上午）', level=1)

doc.add_heading('大盘指数', level=2)
add_styled_table(doc,
    ['指数', '收盘', '涨跌幅'],
    [
        ['上证指数', '3927.60', '-0.53%'],
        ['深证成指', '13548.46', '-1.15%'],
        ['创业板指', '3189.72', '-1.78%'],
        ['科创50', '1268.90', '-2.26%'],
    ])
doc.add_paragraph('')
doc.add_paragraph('热点：医药板块集体爆发，欧康医药+30%涨停，星辉环材+20%涨停。')
doc.add_paragraph('弱势：光模块/AI硬件回调，阳光电源、中际旭创大资金流出。')

# ============ 第二部分 ============
doc.add_heading('二、002788 鹭燕医药 深度分析', level=1)

doc.add_heading('2.1 近期走势', level=2)
add_styled_table(doc,
    ['日期', '收盘', '涨跌幅', '成交额', '换手率'],
    [
        ['3/30', '14.70', '+7.22%', '8.9亿', '16.01%'],
        ['3/31', '16.17', '+10.00%涨停', '16.3亿', '27.32%'],
        ['4/1', '16.70', '+3.28%', '19.0亿', '30.91%'],
        ['4/2上午', '17.10', '+2.40%', '11.3亿', '17.77%'],
    ])

doc.add_heading('2.2 3/31分析回顾与验证', level=2)
doc.add_paragraph('原始判断：底部启动形态，涨停板质量好，建仓位16.0', style='List Bullet')
doc.add_paragraph('验证结果：✅ 4天+16.3%完全验证，启动力度比预期更强', style='List Bullet')

doc.add_heading('2.3 今日千档盘口（4/2 13:05）', level=2)
add_styled_table(doc,
    ['档位', '卖价', '卖量(手)', '买价', '买量(手)'],
    [
        ['5档', '17.17', '56', '17.11', '30'],
        ['4档', '17.16', '48', '17.10', '144'],
        ['3档', '17.15', '86', '17.09', '35'],
        ['2档', '17.14', '57', '17.08', '39'],
        ['1档', '17.13', '49', '17.07', '18'],
    ])
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('盘口意图：').bold = True
doc.add_paragraph('卖盘极轻（五档合计仅296手），筹码高度锁定', style='List Bullet')
doc.add_paragraph('买3(17.10)挂144手防守，是多头底线', style='List Bullet')
doc.add_paragraph('盘口极轻=变盘前兆，买卖都在观望=分歧中=买点', style='List Bullet')

doc.add_heading('2.4 关键价位', level=2)
add_styled_table(doc,
    ['位置', '价位', '意义'],
    [
        ['强支撑', '15.71', '今日最低，破了就危险'],
        ['支撑', '16.00-16.08', '建仓区+开盘价'],
        ['多空分界', '17.00-17.10', '午收价+买3防守位'],
        ['阻力/强阻力', '17.41-18.00', '今日最高+心理关口'],
    ])

# ============ 第三部分 ============
doc.add_heading('三、600666 奥瑞德 分析', level=1)
doc.add_paragraph('节奏：拉升→试盘→洗盘→再拉升，周期2-3天', style='List Bullet')
doc.add_paragraph('4/1涨停6.27，4/2待确认走势', style='List Bullet')
doc.add_paragraph('主力成本区：5.50-6.00 | 关键支撑：6.00-6.10', style='List Bullet')

# ============ 第四部分 ============
doc.add_heading('四、核心交易理论', level=1)
p = doc.add_paragraph()
p.add_run('买在分歧，卖在人声鼎沸处').bold = True
doc.add_paragraph('当前状态：盘口极轻，买卖观望，没人声没鼎沸 = 买点', style='List Bullet')
doc.add_paragraph('卖出信号：成交额爆25亿+，盘口挂满买单，散户追高', style='List Bullet')

p = doc.add_paragraph()
p.add_run('买卖逻辑一致性').bold = True
doc.add_paragraph('支撑位=观察企稳后买入（跟随）', style='List Bullet')
doc.add_paragraph('阻力位=观察是否突破，而非预判顶部（跟随）', style='List Bullet')

# ============ 第五部分 ============
doc.add_heading('五、仓位分配方案', level=1)
add_styled_table(doc,
    ['项目', '002788 鹭燕医药', '600666 奥瑞德', '预留'],
    [
        ['分配比例', '50%', '30%', '20%'],
        ['金额', '6,100元', '3,660元', '2,440元'],
        ['目的', '主仓（趋势确定性高）', '副仓（博洗盘后拉升）', '补仓/应急'],
    ])

# ============ 第六部分 ============
doc.add_heading('六、002788 建仓计划', level=1)
add_styled_table(doc,
    ['步骤', '条件', '仓位', '金额/股数', '操作'],
    [
        ['试盘', '回调17.00-17.10', '1成', '~1200元(约700股)', '小仓试探'],
        ['加仓', '突破17.41放量', '+1成', '~1200元', '确认突破后追'],
        ['满仓', '站稳18.00', '+1成', '~1200元', '主升浪确认'],
        ['止损', '跌破15.70', '全出', '—', '无条件执行'],
    ])

doc.add_heading('七、600666 建仓计划', level=1)
add_styled_table(doc,
    ['步骤', '条件', '仓位', '金额/股数', '操作'],
    [
        ['等待', '看今日收盘定支撑', '—', '—', '不急'],
        ['建仓', '回踩6.00-6.10支撑', '1成', '~1200元(约2000股)', '成本区附近'],
        ['加仓', '放量突破6.65前高', '+1成', '~1200元', '突破确认'],
        ['止损', '跌破5.70', '全出', '—', '前低下方'],
    ])

# ============ 第八部分 ============
doc.add_heading('八、风险控制纪律', level=1)
doc.add_paragraph('单只票最大亏损不超过总资金5%（610元）', style='List Bullet')
doc.add_paragraph('总仓位不超过80%，预留20%应对突发', style='List Bullet')
doc.add_paragraph('止损无条件执行，不扛单不补仓', style='List Bullet')
doc.add_paragraph('先做002788一只，600666等信号再动手（可选）', style='List Bullet')

# ============ 第九部分 ============
doc.add_heading('九、盈亏比分析', level=1)

doc.add_heading('002788（以17.10买入计）', level=2)
add_styled_table(doc,
    ['目标/止损', '价位', '幅度', '盈亏比'],
    [
        ['止损', '15.70', '-8.2%', '—'],
        ['目标1', '18.50', '+8.2%', '1:1'],
        ['目标2', '20.00', '+17.0%', '1:2.1'],
    ])

doc.add_heading('600666（以6.10买入计）', level=2)
add_styled_table(doc,
    ['目标/止损', '价位', '幅度', '盈亏比'],
    [
        ['止损', '5.70', '-6.6%', '—'],
        ['目标1', '6.65', '+9.0%', '1:1.4'],
        ['目标2', '7.50', '+23.0%', '1:3.5'],
    ])

# ============ 免责 ============
doc.add_heading('免责声明', level=1)
doc.add_paragraph('以上分析仅供参考，不构成投资建议。股市有风险，投资需谨慎。')

# 保存
output_path = '/home/adam/.openclaw/workspace/002788_600666_建仓计划_20260402.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
