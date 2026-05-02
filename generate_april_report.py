#!/usr/bin/env python3
"""生成 2026年4月投资建议报告 .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd
import os
from datetime import datetime

# 读取数据
csv_file = '/home/adam/.openclaw/workspace/april_targets_data.csv'
df = pd.read_csv(csv_file)

doc = Document()

# 默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_colored_heading(doc, text, color=RGBColor(139, 0, 0)):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = color
    return p

# ===== 封面 =====
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年4月投资建议')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('霍尔木兹海峡危机下的中国煤化工与油气投资机会')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(60)
run = p.add_run(f'报告日期：2026年3月29日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('⚠️ 本报告仅供参考，不构成投资建议')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(180, 0, 0)

doc.add_page_break()

# ===== 一、核心观点 =====
add_colored_heading(doc, '一、核心观点')

p = doc.add_paragraph()
run = p.add_run('一句话结论：')
run.font.bold = True
run.font.size = Pt(12)
p.add_run(
    '霍尔木兹海峡危机使中国煤化工/油气龙头享受"原料自给+产品涨价"双重红利，'
    '4月仍是最佳配置窗口，但需区分"已涨到位"与"尚有空间"的标的，分批布局。'
)

# 关键判断表
key_views = [
    ['维度', '判断', '说明'],
    ['大方向', '看多 ✅', '危机持续是大概率事件，4月板块仍有上行空间'],
    ['首选标的', '华鲁恒升、宝丰能源', '华鲁恒升3月回调近18%，估值更具吸引力；宝丰能源龙头地位稳固'],
    ['次选标的', '中国石油、中国海油', '油价受益确定性强，防御性好，适合稳健配置'],
    ['警惕标的', '荣盛石化、恒力石化', '炼化一体化受原料成本上涨压制，弹性不如煤化工'],
    ['关键变量', '海峡何时恢复通航', '若全面恢复，板块将面临快速回调'],
    ['建议仓位', '总仓位30-40%', '分两批建仓，上旬一批，中旬视局势加仓'],
]

table = doc.add_table(rows=len(key_views), cols=3)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(key_views):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
                if i == 1 and j == 1:
                    run.font.color.rgb = RGBColor(0, 128, 0)

# ===== 二、市场数据 =====
add_colored_heading(doc, '二、核心标的最新数据（截至3月28日）')

# 数据表
display_cols = ['代码', '名称', '最新收盘', '5日涨幅(%)', '30日涨幅(%)', '年初至今(%)', 
                '距高点回撤(%)', '年化波动率(%)']
available_cols = [c for c in display_cols if c in df.columns]

table = doc.add_table(rows=len(df)+1, cols=len(available_cols))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
for j, header in enumerate(available_cols):
    cell = table.cell(0, j)
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# 数据行
for i, (_, row) in enumerate(df.iterrows()):
    for j, col in enumerate(available_cols):
        val = row[col]
        cell = table.cell(i+1, j)
        if col in ['5日涨幅(%)', '30日涨幅(%)', '年初至今(%)', '距高点回撤(%)']:
            cell.text = f"{val:+.2f}%"
        else:
            cell.text = str(val)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                # 涨幅颜色
                if col in ['5日涨幅(%)', '30日涨幅(%)', '年初至今(%)'] and val > 0:
                    run.font.color.rgb = RGBColor(220, 20, 20)
                elif col in ['5日涨幅(%)', '30日涨幅(%)', '年初至今(%)'] and val < 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)

# ===== 三、投资组合推荐 =====
add_colored_heading(doc, '三、4月投资组合推荐')

doc.add_heading('3.1 核心仓位（60%）— 煤化工龙头', level=2)

p = doc.add_paragraph()
p.add_run('选股逻辑：').bold = True
p.add_run('原料完全自给，产品价格弹性大，危机持续期间利润确定性最高。')

coal_portfolio = [
    ['标的', '代码', '3月涨幅', '距高点', '推荐理由', '建议仓位'],
    ['华鲁恒升', '600426.SH', '-0.90%', '-17.66%', '3月回调充分，估值具吸引力；煤化工多联产龙头，产品多元化分散风险', '20%'],
    ['宝丰能源', '600989.SH', '+26.74%', '-17.51%', '煤制烯烃龙头，成本行业最低；产能扩张中，成长性突出', '20%'],
    ['卫星化学', '002648.SZ', '+20.10%', '-7.40%', 'C2/C3产业链完整；丙烯酸全球龙头，下游需求刚性', '10%'],
    ['中煤能源', '601898.SH', '+19.28%', '-9.56%', '煤炭+煤化工一体化；估值低，股息率高', '10%'],
]

table = doc.add_table(rows=len(coal_portfolio), cols=6)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(coal_portfolio):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.2 防御仓位（25%）— 国内油气开采', level=2)

p = doc.add_paragraph()
p.add_run('选股逻辑：').bold = True
p.add_run('油价直接受益，业绩确定性强，波动率相对较低，提供组合防御性。')

oil_portfolio = [
    ['标的', '代码', '3月涨幅', '距高点', '推荐理由', '建议仓位'],
    ['中国石油', '601857.SH', '+8.84%', '-11.83%', '国内最大油气生产商；高股息+油价弹性；"中特估"概念', '15%'],
    ['中国海油', '600938.SH', '+10.34%', '-7.79%', '纯上游标的，油价弹性最大；海上油气储量丰富', '10%'],
]

table = doc.add_table(rows=len(oil_portfolio), cols=6)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(oil_portfolio):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.3 机动仓位（15%）— 波段交易', level=2)

p = doc.add_paragraph()
p.add_run('选股逻辑：').bold = True
p.add_run('高弹性标的，适合短线波段操作，需严格止损。')

tactical = [
    ['标的', '代码', '特征', '操作建议'],
    ['宝泰隆', '601011.SH', '小盘煤化工，弹性大，波动剧烈', '回调5日线附近买入，涨15%止盈'],
    ['海油工程', '600583.SH', '油气工程，受益于上游资本开支增加', '突破20日高点追入，跌破前低止损'],
    ['中海油服', '601808.SH', '海上油服龙头，订单增长预期', '类似海油工程策略'],
]

table = doc.add_table(rows=len(tactical), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(tactical):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== 四、建仓策略 =====
add_colored_heading(doc, '四、建仓策略与节奏')

doc.add_heading('4.1 分批建仓计划', level=2)

build_plan = [
    ['批次', '时间', '仓位', '操作', '条件'],
    ['第一批', '4月1-3日', '总仓位50%', '按目标仓位50%建仓', '节后开盘直接执行'],
    ['第二批', '4月7-11日', '总仓位30%', '加仓至80%', '若海峡局势未缓和，确认后加仓'],
    ['预留', '4月中下旬', '20%现金', '视情况加仓或防御', '若冲突升级满仓；若缓和则保持现金'],
]

table = doc.add_table(rows=len(build_plan), cols=5)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(build_plan):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.2 买入价位参考', level=2)

p = doc.add_paragraph('基于3月28日收盘价，建议的买入价位：')

entry_prices = [
    ['标的', '当前价', '建议买入价', '止损价', '目标价', '预期收益'],
    ['华鲁恒升', '36.31', '35.00-36.50', '32.00', '45.00', '+22%'],
    ['宝丰能源', '30.10', '28.00-30.00', '25.50', '38.00', '+27%'],
    ['卫星化学', '27.78', '26.50-28.00', '23.50', '34.00', '+22%'],
    ['中煤能源', '17.88', '16.80-18.00', '15.00', '22.00', '+23%'],
    ['中国石油', '12.07', '11.50-12.20', '10.50', '14.50', '+20%'],
    ['中国海油', '41.07', '38.50-41.00', '35.00', '50.00', '+22%'],
]

table = doc.add_table(rows=len(entry_prices), cols=6)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(entry_prices):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True

# ===== 五、情景应对 =====
add_colored_heading(doc, '五、不同情景下的应对策略')

scenarios = [
    ['情景', '触发信号', '操作策略', '预期效果'],
    ['冲突升级\n（概率：中高）',
     '海峡进一步封锁\n油价突破130美元\n化工品价格加速上涨',
     '满仓持有\n减持防御仓加码煤化工\n适当加杠杆（如融资）',
     '板块涨幅加速\n煤化工利润超预期\n组合收益最大化'],
    ['维持僵局\n（概率：中等）',
     '海峡未恢复但无升级\n油价100-120美元震荡\n化工品高位运行',
     '保持80%仓位\n耐心持有核心标的\n等待一季报催化',
     '板块震荡上行\n时间换空间\n享受业绩兑现'],
    ['部分恢复\n（概率：中低）',
     '有限度恢复通航\n油价回落至85-100美元\n化工品价格回落',
     '减仓至50%\n保留龙头去除弹性标的\n增加现金比例',
     '锁定部分利润\n降低组合波动\n防御为主'],
    ['快速缓和\n（概率：低）',
     '全面恢复通航\n油价快速回落至80美元以下\n市场风险偏好骤降',
     '果断减仓至20%以下\n止损弱势股\n保留最强龙头',
     '保护本金\n减少损失\n等待下一次机会'],
]

table = doc.add_table(rows=len(scenarios), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(scenarios):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== 六、风险控制 =====
add_colored_heading(doc, '六、风险控制要点')

risk_rules = [
    ('止损纪律', '单只标的亏损超过15%无条件止损；组合整体亏损超过10%降至半仓。'),
    ('仓位控制', '任何时候保留至少10%现金，不追涨杀跌，不满仓单押一个方向。'),
    ('信息跟踪', '每日关注：霍尔木兹海峡新闻、国际油价走势、化工品现货价格、一季报预告。'),
    ('止盈纪律', '单只标的盈利超过40%开始分批止盈；板块整体涨幅超过50%降低仓位至50%。'),
    ('时间止损', '若4月中旬板块仍未启动（涨幅<5%），检视逻辑是否成立，必要时减仓。'),
    ('分散持仓', '至少持有3个以上标的，单一标的不超过总仓位25%。'),
]

for title, content in risk_rules:
    p = doc.add_paragraph()
    run = p.add_run(f'【{title}】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(139, 0, 0)
    p.add_run(f' {content}')

# ===== 七、关键跟踪指标 =====
add_colored_heading(doc, '七、4月需重点跟踪的指标')

track_items = [
    ['指标', '当前值', '看多信号', '看空信号'],
    ['霍尔木兹海峡通航情况', '禁航中', '持续封锁/升级', '恢复通航'],
    ['布伦特原油价格', '~120美元/桶', '>100美元/桶', '<85美元/桶'],
    ['化工品现货价格（LLDPE）', '持续上涨', '周涨幅>3%', '周跌幅>5%'],
    ['煤化工开工率', '高位', '>85%', '<70%'],
    ['IEA战略储备释放', '4亿桶已释放', '不再追加释放', '大规模追加释放'],
    ['中美关系动态', '关注', '缓和迹象', '对抗升级'],
    ['一季报预告', '4月中旬开始', '煤化工预增>100%', '低于预期'],
]

table = doc.add_table(rows=len(track_items), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(track_items):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
                if j == 2:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif j == 3:
                    run.font.color.rgb = RGBColor(220, 20, 20)

# ===== 八、风险声明 =====
add_colored_heading(doc, '八、风险声明', RGBColor(100, 100, 100))

disclaimers = [
    '本报告仅为个人观点，不构成任何投资建议或承诺。投资有风险，入市需谨慎。',
    '报告中的数据来源于公开信息，可能存在偏差或滞后。',
    '地缘政治局势瞬息万变，报告中的判断可能因突发事件而失效。',
    '过往业绩不代表未来表现，任何投资决策需结合个人风险承受能力。',
    '投资者应根据自身情况独立做出投资决策，并承担相应风险和损失。',
]
for d in disclaimers:
    doc.add_paragraph(d, style='List Bullet')

# 页脚
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

# 保存
output_file = '/home/adam/.openclaw/workspace/2026年4月投资建议_煤化工油气板块.docx'
doc.save(output_file)
print(f'✅ 报告已生成：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
