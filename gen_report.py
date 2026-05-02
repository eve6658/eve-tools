from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 标题 =====
title = doc.add_heading('', level=0)
run = title.add_run('600666 奥瑞德 — 2026年4月2日全天盘口分析报告')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题信息
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('分析时间：09:25-11:33\n分析师：Eve（AI交易助手）')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # 空行

# ===== 辅助函数 =====
def add_section_title(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_sub_title(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

def add_body(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

# ============================
# 一、集合竞价分析
# ============================
add_section_title('一、集合竞价分析（09:25）')
add_body('开盘价 6.27，平开（=昨收价）')
add_body('解读：集合竞价无抢筹也无恐慌，多空平衡，主力不急于拉升。')

# ============================
# 二、早盘冲高
# ============================
add_section_title('二、早盘冲高（09:31-09:33）')

add_sub_title('分时数据')
add_table(
    ['时间', '开盘', '最新', '最高', '最低', '成交量'],
    [
        ['09:31', '6.27', '6.23', '6.39', '6.17', '75万手'],
        ['09:32', '6.21', '6.46', '6.46', '6.21', '34万手'],
        ['09:33', '6.47', '6.60', '6.65', '6.43', '46万手'],
    ]
)

add_sub_title('三阶段解读')
add_body('1. 09:31 试盘：平开后下探6.17再拉回6.39，试探下方承接力', bold=False)
add_body('2. 09:32-09:33 强攻：2分钟内从6.21拉到6.65（+7.1%），突破3/30高点6.47', bold=False)
add_body('3. 目的：快速脱离主力成本区（5.5-6.0），不给散户上车机会', bold=False)

# ============================
# 三、缩量回落
# ============================
add_section_title('三、缩量回落（09:34-10:14）')

add_sub_title('关键节点')
add_body('09:34 6.46 → 09:39 6.33 → 09:46 6.20 → 09:49 6.12 → 10:14 6.14')
add_body('成交量从46万手递减到每分钟1-2万手')

add_body('解读：缩量回落不是出货，是主力主动压盘控节奏。如果出货应该放量下跌。')

# ============================
# 四、千档盘口分析
# ============================
add_section_title('四、千档盘口分析（10:15）')

add_sub_title('买盘支撑')
add_bullet('买1：6.16 × 4049手')
add_bullet('买2：6.15 × 1.21万手')
add_bullet('买5：6.12 × 1.34万手')
add_bullet('买6：6.10 × 2.15万手（铁底）')
add_bullet('买9：6.08 × 1.37万手')

add_sub_title('卖盘压力')
add_bullet('卖4：6.20 × 1.36万手')
add_bullet('卖13：6.29 × 1.50万手')
add_bullet('卖15：6.31 × 1.91万手（天花板）')

add_body('解读：买盘三层铁壁支撑，卖盘三道墙压制。主力在6.10-6.30区间震荡吸筹。')

# ============================
# 五、逐笔成交分析
# ============================
add_section_title('五、逐笔成交分析（10:21）')

add_sub_title('买卖对比变化')
add_bullet('10:15：卖盘48.41万手 vs 买盘23.58万手（2:1）')
add_bullet('10:21：卖盘46.67万手 vs 买盘24.88万手（改善）')

add_sub_title('关键发现')
add_bullet('卖1的978手由25笔散户小单组成，平均每笔39手')
add_bullet('买1的1669手有1388手大户单，平均每笔152手')
add_body('结论：散户在卖，主力在买。', bold=True)

# ============================
# 六、主力意图综合判断
# ============================
add_section_title('六、主力意图综合判断（10:30）')

add_sub_title('主力不希望现在散户建仓的理由')
add_bullet('对手盘关系：散户进来分利润')
add_bullet('操作痕迹：平开→快速拉升→立刻回落，制造"不敢进"的氛围')
add_bullet('真正需要散户的时刻只有两个：顶部出货和洗盘末期')

add_body('综合结论：主力控盘蓄势，拒绝散户上车。主力成本区5.50-6.00。', bold=True)

# ============================
# 七、散户视角盘口解读
# ============================
add_section_title('七、散户视角盘口解读（10:14）')

add_sub_title('三层心理分析')
add_body('第一层（恐惧）：卖盘48万手太大，要崩')
add_body('第二层（冷静）：买4到买8全是过万手大单，6.10有2.15万手铁底')
add_body('第三层（真相）：散户在卖，主力在买，盘口在恐吓散户交出筹码')

add_sub_title('主力剧本')
add_body('拉升 → 回落套人 → 阴跌恐慌 → 盘面恐吓 → 散户割肉 → 主力接筹 → 择日拉升')

# ============================
# 八、盘口变化验证
# ============================
add_section_title('八、盘口变化验证（10:21）')

add_sub_title('关键变化')
add_bullet('卖盘减少1.74万手（筹码在被吸收）')
add_bullet('买盘增加1.3万手（资金在入场）')
add_bullet('6.10铁底从2.15万手增加到2.38万手（主力加码支撑）')
add_bullet('买1从53手变2131手（散户开始在6.15挂买单）')

# ============================
# 九、盘中数据验证
# ============================
add_section_title('九、盘中数据验证（11:25）')

add_sub_title('千档盘口重大变化')
add_bullet('现价从6.15升到6.20')
add_bullet('买盘从24.88万手暴增到37.33万手（+12.45万手！）')
add_bullet('卖盘从46.67万手减少到37.37万手（-9.3万手！）')
add_bullet('买卖比从2:1缩小到接近1:1')

add_body('解读：有大资金在一小时内吃了9万手卖盘，同时新增12万手买盘，确认资金净流入。')

# ============================
# 十、主力出货判断
# ============================
add_section_title('十、主力出货判断（11:26）')

add_sub_title('出货特征对比')
add_table(
    ['出货特征', '今天实际', '是否符合'],
    [
        ['放量下跌', '缩量下跌', '❌'],
        ['大单主动卖出', '无>500手绿色大单', '❌'],
        ['卖盘快速增加', '卖盘从48万降到37万', '❌'],
        ['价格快速下坠', '6.11反弹到6.20', '❌'],
        ['反弹时放量诱多', '反弹量没放大', '❌'],
    ]
)
doc.add_paragraph()
add_body('结论：主力没有出货，而是用小单悄悄吸筹。目标不在6.65，他要的更多。', bold=True)

# ============================
# 十一、上午收盘
# ============================
add_section_title('十一、上午收盘（11:33）')
add_bullet('收盘价：6.30（+0.48%）')
add_bullet('成交量：560万手，成交额35.42亿')
add_bullet('最高：6.65，最低：6.11')

add_body('验证了全天判断：6.11-6.14是洗盘低点，主力在低位吸筹后拉回6.30收盘。')

# ============================
# 十二、建仓策略建议
# ============================
add_section_title('十二、建仓策略建议')

add_sub_title('方案一：1000股（推荐）')
add_bullet('第一批500股挂6.10-6.12（主力托单区）')
add_bullet('第二批500股等突破6.25-6.30（确认启动）')
add_bullet('止损：6.00以下全部走')

add_sub_title('关键价位')
add_bullet('止损位：6.00')
add_bullet('目标位：突破6.30看6.65，突破6.65看7.00+')

# ============================
# 十三、小单占比研究课题
# ============================
add_section_title('十三、小单占比研究课题')
add_body('记录：主力用小单建仓/出货时小单占比异常升高，这是必要不充分条件。需要结合价格区间稳定性、成交量方向、压单/托单综合判断。后续通过交易复盘案例量化经验阈值。')

# ============================
# 风险提示
# ============================
doc.add_paragraph()
risk = doc.add_paragraph()
risk.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = risk.add_run('⚠️ 风险提示')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_bullet('奥瑞德属于小市值题材股，波动极大')
add_bullet('以上分析基于盘口数据，不构成投资建议')

# ===== 保存 =====
output_path = '/home/adam/.openclaw/workspace/奥瑞德600666_全天分析报告_20260402.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
