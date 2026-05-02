from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            t.rows[r+1].cells[c].text = val

# 封面
doc.add_paragraph('\n\n')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('A股不对称战场\n操盘手册（完整版）')
run.font.size = Pt(28); run.bold = True
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('截至 2026-04-19 | 含18章 + 实盘案例库')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(100,100,100)
doc.add_page_break()

# 目录
doc.add_heading('目录', level=1)
toc = [
    ('第一部分：盘口微观分析', True),
    ('  第一章 工具中性原则', False),
    ('  第二章 小单占比阈值', False),
    ('  第三章 价格墙理论', False),
    ('  第四章 成交量的本质', False),
    ('  第五章 千档盘口分析方法', False),
    ('第二部分：主力行为分析', True),
    ('  第六章 主力操盘三种手法', False),
    ('  第七章 拉高出货识别', False),
    ('  第八章 主力控盘度判断', False),
    ('  第九章 主力拉升动力分析', False),
    ('第三部分：选股与择时', True),
    ('  第十章 股东人数选股', False),
    ('  第十一章 价格区间选股', False),
    ('  第十二章 趋势跟踪三步法', False),
    ('  第十三章 盈亏比与仓位管理', False),
    ('第四部分：操盘手理论（庄家视角）', True),
    ('  第十四章 控盘坐庄模型推演', False),
    ('  第十五章 妖股形态必要条件', False),
    ('第五部分：心智纪律', True),
    ('  第十六章 无我交易', False),
    ('  第十七章 震荡市纪律', False),
    ('  第十八章 周末效应', False),
    ('第六部分：实盘案例库', True),
    ('附录A 交易成本对比', False),
    ('附录B 关键阈值速查表', False),
]
for item, bold in toc:
    p = doc.add_paragraph(item)
    if bold: p.runs[0].bold = True
doc.add_page_break()

# 第一部分
doc.add_heading('第一部分：盘口微观分析', level=1)
doc.add_heading('第一章 工具中性原则', level=2)
doc.add_paragraph('核心理念：技术指标（MACD/RSI/均线）不预判未来，只描述当前状态。')
for s in ['用"如果...那么..."代替"所以...会..."。', '阻力位是观察区，不是卖出点。', '买卖逻辑必须对称：跟随买入，跟随卖出。']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('第二章 小单占比阈值', level=2)
doc.add_paragraph('数据来源：2026-04-02 实盘量化（600666奥瑞德 + 002788鹭燕）')
add_table(doc, ['场景', '阈值条件', '含义'], [
    ['吸筹', '卖盘小单>80% + 买盘小单<30% + 价格稳定 + 铁底未动', '主力在收集筹码'],
    ['出货', '买盘小单>80% + 卖盘大单密集 + 价格区间下移', '主力在派发'],
    ['散户主导', '买卖盘小单占比均衡(各50-60%)，无明显大单', '无主力参与'],
    ['对倒', '买卖盘同时出现大单，金额接近，快速交替', '主力在做量'],
])
doc.add_paragraph()
doc.add_paragraph('实操四步判断法：')
for i, s in enumerate(['看逐笔成交的买卖比', '统计买/卖盘小单占比', '对照价格区间是否稳定', '检查铁底（万手托单）是否存在']):
    doc.add_paragraph(f'第{i+1}步：{s}', style='List Number')

doc.add_heading('第三章 价格墙理论', level=2)
doc.add_paragraph('核心定义：价格墙 = 某价位堆积万手级别挂单（压单=天花板，托单=铁底）。')
doc.add_paragraph('核心机制：控制者在现价上下挂大单，中间区间是散户投票区。控制者通过反向操作收割：')
for s in ['人多 → 撤地板下跌洗盘', '人少 → 撤天花板喷发', '30%止损率：马丁策略+长持者不会被洗，正好洗30%摇摆筹码。']:
    doc.add_paragraph(s, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('价格墙厚度变化规律：')
add_table(doc, ['变化方向', '含义'], [
    ['25万→10万→5万→消失', '主力撤退，准备发车'],
    ['5万→10万→25万', '主力加压，准备出货'],
])

doc.add_heading('第四章 成交量的本质', level=2)
doc.add_paragraph('核心认知：成交量是分歧指标，不是方向信号。放量只能说明买卖双方想法不一样，不知道谁对。')
doc.add_paragraph('四种量价关系：')
add_table(doc, ['信号', '含义'], [
    ['底部无量涨停', '没人卖，真强'],
    ['底部放量涨停', '有人卖，可能对倒，观察次日'],
    ['高位放巨量冲高回落', '主力大规模派发，危险'],
    ['高位缩量下跌', '没人接盘，更危险'],
])

doc.add_heading('第五章 千档盘口分析方法', level=2)
doc.add_paragraph('工具：本地OCR（skills/local-ocr/ocr_logic.py），禁止用视觉模型。')
doc.add_paragraph('关注三要素：委比变化、买/卖盘总量对比、大单分布。')
doc.add_paragraph('核心规则：委比 = 哪边人多 → 人多的那边要被反向收割。')
doc.add_page_break()

# 第二部分
doc.add_heading('第二部分：主力行为分析', level=1)
doc.add_heading('第六章 主力操盘三种手法', level=2)
add_table(doc, ['手法', '模式', '案例'], [
    ['A手法', '拉高→封板→消耗封单→次日低开（单边出货）', '鹭燕'],
    ['B手法', '吸筹→出货→回补→再拉→再出（两头剥皮）', '奥瑞德'],
    ['C手法', '连续涨停→高位横盘出货→阴跌', '妖股模式'],
])

doc.add_heading('第七章 拉高出货识别', level=2)
doc.add_paragraph('完整识别链条（2026-04-16 鹭燕完整验证）：')
for i, s in enumerate(['封单消耗', '集合竞价未涨停', '盘口委比翻负', '超大单净流出', '逐笔卖:买比例恶化', '价格崩塌']):
    doc.add_paragraph(f'第{i+1}步：{s}', style='List Number')
doc.add_paragraph()
doc.add_paragraph('关键信号阈值：')
for s in ['集合竞价未涨停 = 出货确认第一步', '委比从正转负 = 诱多结束', '超大单4分钟净流出 >1000万 = 实锤', '逐笔卖:买 >3:1 = 加速出货']:
    doc.add_paragraph(s, style='List Bullet')
p = doc.add_paragraph()
r = p.add_run('重要纠正（Adam）：出货区间下沿 ≠ 支撑位。'); r.bold = True

doc.add_heading('第八章 主力控盘度判断', level=2)
doc.add_paragraph('控盘甜区：流通盘的30-50%。')
for s in ['80%控盘理论上利润最大，但现实中几乎不可能（举牌线+减持新规限制）。', '50%控盘实操中最灵活，进退有余。']:
    doc.add_paragraph(s, style='List Bullet')
doc.add_paragraph('控盘度判断方法：')
for s in ['股东人数减少 + 人均持股量增加', '成交量萎缩但价格稳定', '极度缩量（日成交<常态的1/3）']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('第九章 主力拉升动力分析', level=2)
doc.add_paragraph('拉升三条件（缺一不可）：')
for s in ['1. 筹码清洗完毕（极度缩量=筹码锁定）', '2. 板块有风（板块轮动主线）', '3. 大盘配合（稳定或上涨）']:
    doc.add_paragraph(s)
doc.add_paragraph()
doc.add_paragraph('拉升历史规律：极度缩量 → 突然放量涨停 → 阴跌出货')
doc.add_paragraph('主力成本估算：横盘区均价 ≈ 主力成本')
doc.add_page_break()

# 第三部分
doc.add_heading('第三部分：选股与择时', level=1)
doc.add_heading('第十章 股东人数选股', level=2)
doc.add_paragraph('甜区：2万-20万人')
add_table(doc, ['人数区间', '含义'], [
    ['2-5万人', '高度集中，主力控盘迹象明显'],
    ['5-10万人', '中等集中，主力建仓中期'],
    ['10-20万人', '开始分散，主力可能在派发'],
    ['>20万人', '散户化，基本没戏'],
    ['<2万人', '筹码太集中或僵尸股'],
])
doc.add_paragraph()
doc.add_paragraph('披露频率差异：')
doc.add_paragraph('深交所（0开头）：每10个交易日更新一次（一个月2-3次）', style='List Bullet')
doc.add_paragraph('上交所（6开头）：只在季报/年报披露（一个季度1次）', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('最佳信号：股东人数连续2-3期递减 + 人均持股量上升 + 股价低位横盘')
doc.add_paragraph('⚠️ 数据有10天滞后，需配合量价关系判断')

doc.add_heading('第十一章 价格区间选股', level=2)
doc.add_paragraph('散户最佳机会窗口：10-50元')
doc.add_paragraph('5元以下：低价股，流动性差，盈亏比低，主力不愿碰', style='List Bullet')
doc.add_paragraph('100元以上：不是散户机会，看都别看', style='List Bullet')

doc.add_heading('第十二章 趋势跟踪三步法', level=2)
doc.add_paragraph('口诀：观察资金在哪 → 等待突破 → 发车再入，倒车离场')
doc.add_paragraph('三不原则：')
for s in ['不猜底（等资金进场再跟）', '不猜顶（看到倒车就走）', '不贪心（不赚最后一个铜板）']:
    doc.add_paragraph(s, style='List Bullet')
p = doc.add_paragraph()
r = p.add_run('铁律：涨停开盘必走。涨停 = 流动性最好的出货窗口。'); r.bold = True

doc.add_heading('第十三章 盈亏比与仓位管理', level=2)
doc.add_paragraph('最优仓位公式（凯利公式）：')
doc.add_paragraph('仓位 = (胜率 × 赔率 - 败率) / 赔率')
doc.add_paragraph('实际使用取半（保守原则）。')
doc.add_paragraph('盈亏比不达标（<2:1）不进场。')
doc.add_paragraph('高盈亏比只在恐惧和不耐烦中产生。')
doc.add_page_break()

# 第四部分
doc.add_heading('第四部分：操盘手理论（庄家视角）', level=1)
doc.add_heading('第十四章 控盘坐庄模型推演', level=2)
doc.add_paragraph('3-15-10形态（经典妖股套路）：')
for s in ['3→5：试盘+吸筹（缓慢放量）', '5→10：主升浪（连续涨停，制造人气）', '10→15：加速赶顶（最疯狂，吸引散户接力）', '15→10：高位派发（横盘震荡出货）']:
    doc.add_paragraph(s, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('50%流通控盘是实操甜区：')
for s in ['外面50%筹码中，真正活跃的只有15-20%', '实际浮动筹码只有总股本的7-10%', '用1-2亿就能拉动股价']:
    doc.add_paragraph(s)
doc.add_paragraph()
doc.add_paragraph('大资金交易成本对比：')
add_table(doc, ['项目', '大资金', '散户'], [
    ['佣金', '万1', '万2.5'],
    ['过户费', '万0.5', '万0.5'],
    ['印花税', '千1（卖出）', '千1（卖出）'],
    ['一个来回', '~千1.3', '~千3.5'],
])
doc.add_paragraph()
doc.add_paragraph('庄家利润三来源：筹码优势 + 成本优势 + 信息优势')

doc.add_heading('第十五章 妖股形态必要条件', level=2)
doc.add_paragraph('同时满足四条才能走出妖股形态：')
for s in ['1. 绝对控盘（>60%，实操中30-50%更常见）', '2. 没有大资金对手盘', '3. 市场情绪配合（牛市或题材热点）', '4. 出货阶段有足够的散户涌入']:
    doc.add_paragraph(s)
doc.add_paragraph()
doc.add_paragraph('四条缺一条，形态就走不出来。看到的妖股都是幸存者偏差——走不出来的K线上根本不会出现3-15-10。')
doc.add_page_break()

# 第五部分
doc.add_heading('第五部分：心智纪律', level=1)
doc.add_heading('第十六章 无我交易', level=2)
doc.add_paragraph('核心测试：如果今天空仓，我会以当前价买入吗？')
doc.add_paragraph('答案是否定的 → 不要因为持仓而持有')
doc.add_paragraph('沉没成本谬误：已亏的钱不是钱，未来的盈亏才是决策依据')

doc.add_heading('第十七章 震荡市纪律', level=2)
doc.add_paragraph('追跌杀涨（不是追涨杀跌）')
doc.add_paragraph('恐惧时买入，兴奋时卖出')

doc.add_heading('第十八章 周末效应', level=2)
doc.add_paragraph('周末前游资通常减仓避险')
doc.add_paragraph('周五不追高，除非板块确定性强')
doc.add_page_break()

# 第六部分
doc.add_heading('第六部分：实盘案例库', level=1)
doc.add_heading('成功案例', level=2)
add_table(doc, ['股票', '日期', '模式', '关键信号', '结果'], [
    ['600666奥瑞德', '2026-04-02', '压盘吸筹', '卖盘小单>80%，买盘平均152手', '4.5建仓，6.5清仓'],
    ['002788鹭燕', '2026-04-03', '试仓卖出', '天量冲高回落+换手30%+', '16.97买200股，17.37卖'],
    ['600666奥瑞德', '2026-04-16', '涨停出货预判', '封单消耗+集合竞价未涨停', '预判完全兑现，14:32清仓'],
])

doc.add_heading('教训案例', level=2)
add_table(doc, ['股票', '问题', '教训'], [
    ['002788鹭燕', '15.70误判为支撑', '出货区间下沿≠支撑位'],
    ['600666奥瑞德', '涨停没走', '涨停开盘=出货窗口，先走不犹豫'],
])
doc.add_page_break()

# 附录
doc.add_heading('附录A：交易成本对比', level=2)
doc.add_paragraph('详见第十四章表格。大资金一个来回约千1.3，散户约千3.5。')

doc.add_heading('附录B：关键阈值速查表', level=2)
add_table(doc, ['指标', '阈值', '含义'], [
    ['股东人数', '2-20万', '可炒作区间'],
    ['股东人数连续下降', '≥3期', '筹码集中化'],
    ['卖盘小单占比', '>80%', '吸筹信号'],
    ['买盘小单占比', '<30%', '吸筹信号'],
    ['换手率', '>25%', '游资接力/高风险'],
    ['换手率', '<5%', '极度缩量/筹码锁定'],
    ['流通盘控盘', '30-50%', '实操甜区'],
    ['最佳股价区间', '10-50元', '散户机会窗口'],
    ['凯利仓位', '取半', '保守原则'],
    ['盈亏比最低线', '≥2:1', '不达标不进场'],
])

out = '/home/adam/.openclaw/workspace/A股不对称战场操盘手册_完整版_2026-04-19.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)} bytes')
