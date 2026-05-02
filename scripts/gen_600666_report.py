#!/usr/bin/env python3
"""生成600666奥瑞德分析报告docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3

# ===== 标题 =====
title = doc.add_heading('600666 奥瑞德 — 主力动向分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'分析日期：2026年3月30日 | 实验建仓价：5.48元')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ===== 一、公司概况 =====
doc.add_heading('一、公司概况', level=1)
p = doc.add_paragraph()
p.add_run('股票代码：').bold = True
p.add_run('600666.SH\n')
p.add_run('公司名称：').bold = True
p.add_run('奥瑞德光电股份有限公司\n')
p.add_run('主营业务：').bold = True
p.add_run('蓝宝石材料生产销售 + 算力租赁业务\n')
p.add_run('第一大股东：').bold = True
p.add_run('青岛智算信息产业发展合伙企业（有限合伙），持股13.02%')

# ===== 二、近期行情回顾 =====
doc.add_heading('二、近期行情回顾（上周至今日）', level=1)

doc.add_heading('2.1 连板拉升阶段（3月24日-27日）', level=2)
p = doc.add_paragraph('上周走出一波强势连板行情：')

table_data = [
    ['日期', '收盘价', '涨跌幅', '成交额(亿)', '量比', '特征'],
    ['3/24', '4.80', '+10.09%', '17.63', '0.95', '涨停启动'],
    ['3/25', '5.28', '+10.00%', '3.47', '0.18', '一字板锁仓'],
    ['3/26', '5.53', '+4.73%', '58.93', '2.00', '放量冲高'],
    ['3/27', '6.08', '+9.95%', '61.20', '1.74', '再封涨停'],
]
table = doc.add_table(rows=len(table_data), cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(table_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
        for paragraph in table.rows[i].cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_heading('2.2 今日走势（3月30日）', level=2)

today_data = [
    ['开盘', '最高', '最低', '收盘', '涨跌幅', '成交额', '换手率'],
    ['5.80', '6.19', '5.47', '5.51→5.97', '-1.81%', '34.69亿', '24.93%'],
]
table2 = doc.add_table(rows=2, cols=7)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(today_data):
    for j, cell_text in enumerate(row_data):
        table2.rows[i].cells[j].text = cell_text
        for paragraph in table2.rows[i].cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('\n今日走势特征：').bold = True
p.add_run('开盘5.80，盘中探底5.47（接近实验建仓价5.48），随后反弹至最高6.19，最终回落至5.97附近。全天振幅13.1%，走"探底反弹再回落"形态。')

# ===== 三、技术指标分析 =====
doc.add_heading('三、技术指标分析', level=1)

tech_data = [
    ['指标', '数值', '判断'],
    ['MA5', '5.44', '价格在MA5上方 ✅'],
    ['MA10', '4.93', '多头排列 ✅'],
    ['MA20', '4.49', '价格远高于MA20，乖离22.62% ⚠️'],
    ['MA60', '3.93', '长期均线多头排列 ✅'],
    ['MACD DIF', '0.4421', 'DIF>DEA，金叉状态 ✅'],
    ['MACD柱', '0.2885', '多头动能 ✅'],
    ['量比', '0.59', '缩量（相比前几日放量后回落）'],
    ['乖离率', '22.62%', '偏高，有回调压力 ⚠️'],
]
table3 = doc.add_table(rows=len(tech_data), cols=3)
table3.style = 'Light Grid Accent 1'
for i, row_data in enumerate(tech_data):
    for j, cell_text in enumerate(row_data):
        table3.rows[i].cells[j].text = cell_text
        for paragraph in table3.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('\n均线排列：').bold = True
p.add_run('MA5>MA10>MA20>MA60，标准多头排列，中长期趋势向好。')

p = doc.add_paragraph()
p.add_run('乖离率偏高：').bold = True
p.add_run('价格离MA20乖离22.62%，说明短期涨幅过大，有技术性回调需求。')

# ===== 四、资金流向分析 =====
doc.add_heading('四、主力资金流向分析', level=1)

doc.add_heading('4.1 近5日主力净流入', level=2)

flow_data = [
    ['日期', '主力净流入(万)', '主力净占比', '特征'],
    ['3/23', '-5,119.22', '-2.50%', '主力撤退'],
    ['3/24', '+33,994.79', '+19.28%', '主力强势买入（涨停日）'],
    ['3/25', '+10,124.36', '+29.16%', '主力锁仓（一字板）'],
    ['3/26', '-53,849.80', '-9.14%', '主力大幅出货 ⚠️'],
    ['3/27', '+35,761.10', '+5.84%', '超大单回流（涨停日）'],
    ['5日合计', '+20,911.23', '—', '整体净流入'],
]
table4 = doc.add_table(rows=len(flow_data), cols=4)
table4.style = 'Light Grid Accent 1'
for i, row_data in enumerate(flow_data):
    for j, cell_text in enumerate(row_data):
        table4.rows[i].cells[j].text = cell_text
        for paragraph in table4.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_heading('4.2 3月27日资金结构拆解', level=2)

struct_data = [
    ['类型', '净流入(万)', '判断'],
    ['超大单', '+37,963.39', '大资金买入 🟢'],
    ['大单', '-2,202.29', '小幅流出'],
    ['中单', '-17,164.53', '中等资金出货 🔴'],
    ['小单', '-18,596.58', '散户出货 🔴'],
]
table5 = doc.add_table(rows=len(struct_data), cols=3)
table5.style = 'Light Grid Accent 1'
for i, row_data in enumerate(struct_data):
    for j, cell_text in enumerate(row_data):
        table5.rows[i].cells[j].text = cell_text
        for paragraph in table5.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('\n资金解读：').bold = True
p.add_run('3/27的主力流入主要来自超大单（+3.8亿），中单和小单均为流出。说明大资金在涨停日接力买入，但散户和中等资金在出货。3/26主力大幅流出5.38亿后，3/27回补3.6亿，主力内部出现明显分歧。')

# ===== 五、主力动向综合判断 =====
doc.add_heading('五、主力动向综合判断', level=1)

p = doc.add_paragraph()
run = p.add_run('核心结论：高位博弈后期，主力分歧加大，短期震荡整理')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(200, 50, 50)

doc.add_heading('5.1 主力行为分析', level=2)

points = [
    ('连板阶段（3/24-27）', '主力快速拉升，换手充分。3/25一字板说明筹码锁定良好，3/26-27放量博弈激烈。'),
    ('3/26主力出货', '单日净流出5.38亿（占比-9.14%），这是明确的出货信号。'),
    ('3/27主力回补', '超大单回流3.8亿封涨停，可能是新主力接力，也可能是原有主力护盘。'),
    ('今日（3/30）', '近跌停开盘后反弹，缩量说明卖盘枯竭但买盘也犹豫。主力分歧尚未分出胜负。'),
]
for title_text, desc in points:
    p = doc.add_paragraph()
    p.add_run(f'• {title_text}：').bold = True
    p.add_run(desc)

doc.add_heading('5.2 综合信号评分', level=2)

score_data = [
    ['维度', '信号', '评分'],
    ['均线系统', '多头排列（MA5>MA10>MA20>MA60）', '✅ 强势'],
    ['MACD', '金叉状态，DIF>DEA', '✅ 多头'],
    ['乖离率', '22.62%，偏高', '⚠️ 回调压力'],
    ['量能', '缩量（量比0.59）', '➡️ 观望'],
    ['主力资金', '5日净流入+2.09亿，但分歧大', '⚠️ 中性偏正'],
    ['盘面形态', '高位震荡，探底反弹再回落', '⚠️ 不确定'],
]
table6 = doc.add_table(rows=len(score_data), cols=3)
table6.style = 'Light Grid Accent 1'
for i, row_data in enumerate(score_data):
    for j, cell_text in enumerate(row_data):
        table6.rows[i].cells[j].text = cell_text
        for paragraph in table6.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# ===== 六、实验建仓分析 =====
doc.add_heading('六、实验建仓分析', level=1)

p = doc.add_paragraph()
p.add_run('建仓价：').bold = True
p.add_run('5.48元（接近今日最低价5.47）\n')
p.add_run('建仓时点：').bold = True
p.add_run('今日盘中低位\n')
p.add_run('当前浮盈：').bold = True
run = p.add_run('约+8.9%（按5.97计）')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True
p.add_run('\n盘中最高浮盈：').bold = True
run = p.add_run('约+13.1%（按6.19高点计）')
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_heading('6.1 关键价位参考', level=2)

level_data = [
    ['价位', '意义', '操作建议'],
    ['5.44', 'MA5支撑位', '跌破则考虑止损'],
    ['5.48', '建仓成本价', '盈亏平衡线'],
    ['5.97', '当前价', '—'],
    ['6.19', '今日最高价', '短线压力位'],
    ['6.08', '3/27涨停收盘价', '近期高点压力'],
]
table7 = doc.add_table(rows=len(level_data), cols=3)
table7.style = 'Light Grid Accent 1'
for i, row_data in enumerate(level_data):
    for j, cell_text in enumerate(row_data):
        table7.rows[i].cells[j].text = cell_text
        for paragraph in table7.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_heading('6.2 操作思路', level=2)

ops = [
    '短线实验仓（当前策略）：+9%利润已可观，可考虑择机减仓锁利',
    '趋势持有：若MA5(5.44)不破，可继续持有观察',
    '止盈位：如再次冲击6.08-6.20区域可考虑止盈',
    '止损位：收盘跌破MA5(5.44)则执行止损',
]
for i, op in enumerate(ops):
    p = doc.add_paragraph(f'{i+1}. {op}')

# ===== 七、风险提示 =====
doc.add_heading('七、风险提示', level=1)

risks = [
    '奥瑞德属于高位连板后回调股，波动极大，不适合重仓',
    '换手率24.93%说明筹码交换充分，但获利盘仍多',
    '主力内部出现分歧，方向尚不明朗',
    '乖离率偏高，技术面有回调至MA20(4.49)的可能',
    '本报告仅为技术面分析，不构成投资建议',
]
for risk in risks:
    doc.add_paragraph(risk, style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n— 报告完 —')
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.size = Pt(9)

# 保存
output_path = '/home/adam/.openclaw/workspace/600666_奥瑞德分析报告_20260330.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
