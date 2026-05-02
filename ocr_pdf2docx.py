#!/usr/bin/env python3
"""OCR a scanned PDF and convert to Word document."""
import sys
import os
import fitz  # PyMuPDF
import pytesseract
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import time

# Setup paths
TESSERACT_BIN = "/home/adam/.local/tesseract/usr/bin/tesseract"
TESSDATA = "/home/adam/.local/tesseract/usr/share/tesseract-ocr/5/tessdata"
LIB_PATH = "/home/adam/.local/tesseract/usr/lib/x86_64-linux-gnu"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_BIN
os.environ['TESSDATA_PREFIX'] = TESSDATA
os.environ['LD_LIBRARY_PATH'] = LIB_PATH + ":" + os.environ.get('LD_LIBRARY_PATH', '')

INPUT_PDF = "/home/adam/文档/1.pdf"
OUTPUT_DOCX = "/home/adam/文档/1.docx"

def ocr_pdf(pdf_path, docx_path):
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    print(f"Total pages: {total_pages}")
    
    word_doc = Document()
    
    # Add title
    word_doc.add_heading('OCR转换结果', level=0)
    word_doc.save(docx_path)  # Save initial
    
    start_time = time.time()
    
    for i, page in enumerate(doc):
        page_start = time.time()
        
        # Render page as high-res image (300 DPI)
        mat = fitz.Matrix(300/72, 300/72)
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PIL Image
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        
        # OCR with Chinese + English
        text = pytesseract.image_to_string(img, lang='chi_sim+eng', config='--psm 6')
        
        # Add page heading
        word_doc.add_heading(f'第 {i+1} 页', level=2)
        
        # Add text
        if text.strip():
            word_doc.add_paragraph(text.strip())
        else:
            word_doc.add_paragraph('[该页无识别到文字]')
        
        # Save progress every 10 pages
        if (i + 1) % 10 == 0:
            word_doc.save(docx_path)
        
        elapsed = time.time() - page_start
        total_elapsed = time.time() - start_time
        avg_time = total_elapsed / (i + 1)
        eta = avg_time * (total_pages - i - 1)
        
        print(f"Page {i+1}/{total_pages} done ({elapsed:.1f}s) | ETA: {eta/60:.1f}min")
    
    word_doc.save(docx_path)
    total_time = time.time() - start_time
    print(f"\nDone! Saved to {docx_path}")
    print(f"Total time: {total_time/60:.1f} minutes")

if __name__ == "__main__":
    ocr_pdf(INPUT_PDF, OUTPUT_DOCX)
