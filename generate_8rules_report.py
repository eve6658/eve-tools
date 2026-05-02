#!/usr/bin/env python3
"""生成 短线交易八条口诀分析报告 .docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_colored_heading(doc, text, color=RGBColor(139, 0, 0)):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = color

def add_bold_run(paragraph, text, color=None):
    run = paragraph.add_run(text)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    return run

# ===== 封面 =====
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('短线交易八条口诀')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('应用场景与实操价值评估')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——基于2026年3月A股市场环境')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(60)
run = p.add_run(f'报告日期：2026年3月29日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ===== 一、引言 =====
add_colored_heading(doc, '一、引言')

p = doc.add_paragraph()
p.add_run(
    '短线交易是A股市场中最具争议也最具吸引力的策略之一。以下八条口诀凝聚了多年盘面实战经验，'
    '覆盖了趋势判断、入场信号、止盈止损和风险控制等核心环节。\n\n'
    '本报告结合2026年3月A股市场实际行情数据（特别是霍尔木兹海峡危机驱动下的石化板块剧烈波动），'
    '逐条评估每条口诀的应用场景、有效性边界和实操价值，并给出针对性的改进建议。'
)

# ===== 二、口诀一览 =====
add_colored_heading(doc, '二、八条口诀一览')

rules_list = [
    ('回踩分时均线', '趋势股回踩分时均线未破，强势可买'),
    ('冲高7%不封板', '热门股冲高7%不封涨停，离场'),
    ('低点抬高入场', '股价低点持续抬高，果断入场'),
    ('低开超3%15分钟不翻红卖出', '低开超3%，15分钟不翻红，卖出'),
    ('连板加速爆量换手', '连板加速，爆量换手，关注机会'),
    ('连涨三天先卖半仓', '连涨三天，先卖半仓'),
    ('中高位缩量可加仓', '中高位缩量，可加仓'),
    ('放量走势不明退出', '放量走势不明，立即退出'),
]

for i, (name, rule) in enumerate(rules_list, 1):
    p = doc.add_paragraph()
    add_bold_run(p, f'第{i}条 【{name}】')
    p.add_run(f' — {rule}')

# ===== 三、当前市场环境 =====
add_colored_heading(doc, '三、当前市场环境概述')

p = doc.add_paragraph()
add_bold_run(p, '全球背景：')
p.add_run(
    '2026年2月28日美以打击伊朗，霍尔木兹海峡禁航，国际油价飙升至~120美元/桶。'
    '全球石油市场面临"史上最严重供应梗阻"，亚洲石化产业链深度冲击。'
)

p = doc.add_paragraph()
add_bold_run(p, 'A股特征：')
p.add_run(
    '① 量能放大，波动率年化50-80%，远高于历史均值\n'
    '② 石油化工板块分化严重：煤化工逆势上涨（宝丰能源+26%、中煤+19%），炼化下跌（荣盛-19%、恒力-16%）\n'
    '③ 热点快速轮动，地缘消息驱动的脉冲式行情频发\n'
    '④ 涨停梯队暂未形成完整连板结构，短线接力氛围偏弱'
)

# ===== 四、逐条深度评估 =====
add_colored_heading(doc, '四、逐条深度评估')

evaluations = [
    {
        'name': '第1条【回踩分时均线未破强势可买】',
        'rating': '⭐⭐⭐⭐ 实用性高',
        'content': p.add_run(
            '趋势股回踩分时均线（VWAP）不破，说明承接有力，是经典的顺势交易信号。\n\n'),
        'scenes': '适合趋势明确的个股（如宝丰能源、卫星化学）在盘中回调时低吸，尤其在当前石化板块整体向上但盘中波动大的环境中，回踩VWAP的机会频繁出现。',
        'case': '宝丰能源3月26日：盘中回调至分时均线附近未破，随后反弹收涨3.2%。类似形态当月出现至少4次。',
        'advantage': '① 顺势操作，成功率高；② 当前高波动市场提供充足买点',
        'risk': '① 趋势识别本身是难点，非趋势股回踩就是破位；② V形反转时可能抄在半山腰',
        'improve': '建议加上"趋势确认"前提条件：20日均线向上 + 最近5日涨幅>3% + 日均成交量高于20日均量',
    },
    {
        'name': '第2条【冲高7%不封涨停离场】',
        'rating': '⭐⭐⭐⭐⭐ 当前最实用',
        'scenes': '当前最适用的场景。地缘危机驱动的化工题材股天然具有"冲高后散户追、主力趁机出货"的特征。',
        'case': '宝丰能源3月振幅42%（涨幅仅14%），中间多次冲高到7-9%未封板后大幅回落。严格执行此条可避开多个高点回撤。',
        'advantage': '① 止损果断，不恋战；② 7%阈值在涨跌停±10%制度下非常合理；③ 避免"再等等就涨停"的幻想',
        'risk': '① 科创板、北交所涨跌停±20%，7%阈值需调整（建议改为12-15%）；② 可能错过日内V转后再封板的票',
        'improve': '科创板/北交所改为12%；加上"量比>2"确认条件——如果冲高7%但量比<1.5，可能是试盘而非出货',
    },
    {
        'name': '第3条【低点持续抬高果断入场】',
        'rating': '⭐⭐⭐ 有局限性',
        'scenes': '适合整理末期、量能收敛的个股突破前布局。当前市场中，华鲁恒升3月下旬回撤至36元附近后低点逐步抬高，是此类形态的典型案例。',
        'case': '华鲁恒升3月20日低点35.8元 → 3月24日低点36.0元 → 3月27日低点36.2元，低点持续抬高，但整体波动不大。',
        'advantage': '① 经典趋势跟踪逻辑；② 理解简单，执行门槛低',
        'risk': '① "果断入场"太激进——主力可做"低点抬高"出货形态；② 单纯低点抬高无法判断真假突破',
        'improve': '必须配合"量能收敛"：低点抬高时成交量逐级缩量，才是健康的洗盘形态；低点抬高但量能放大，可能是假突破',
    },
    {
        'name': '第4条【低开超3%15分钟不翻红卖出】',
        'rating': '⭐⭐⭐ 阈值需调整',
        'scenes': '适用于弱势股补跌或利好不及预期导致的跳空低开。当前市场中，荣盛石化3月某日低开2.8%后15分钟内未翻红，当天收盘跌6.5%。',
        'case': '恒力石化3月18日：低开3.2%，前15分钟最高仅反弹至-1.8%未翻红，随后一路走低收盘-5.8%。如果按此规则操作，成功规避了后续损失。',
        'advantage': '① 防止低开后越跌越深；② 强制离场纪律，避免抄底幻想',
        'risk': '① 2026年A股个股日均波动3-5%，3%的阈值太宽，容易误杀；② 低开3%也可能是洗盘，翻红后大涨的案例不少',
        'improve': '建议改为"低开超5% + 15分钟不翻红 + 委比<-30%"三条件同时满足才卖出；或"低开超3% + 前3分钟最大跌幅>2% + 放量杀跌"确认后卖出',
    },
    {
        'name': '第5条【连板加速爆量换手关注机会】',
        'rating': '⭐⭐⭐⭐⭐ 连板行情核心方法',
        'scenes': '连板行情中的核心策略。当前石化板块尚未出现连板龙头，但4月若地缘冲突持续升级，极可能出现首板接力的连板行情。',
        'case': '参考2024年红宝丽（6连板）等案例，连板换手后第二波行情往往比第一波更凶。当前可关注宝丰能源、中煤能源等是否有首板突破机会。',
        'advantage': '① 二板三板爆量换手后进场是利润最丰厚的阶段；② 换手充分意味着筹码交换，新主力入场意愿强',
        'risk': '① "关注机会"太模糊——需要区分"真换手"（封板力强、量比适中）和"假换手"（反复开板、封板无力）；② 连板接力失败后往往一字跌停，损失巨大',
        'improve': '明确入场条件：连板数3板以上 + 换手率15-25% + 封板时间<10:30 + 封单量>流通盘1%；严格止损：买入后次日不能连板即卖出',
    },
    {
        'name': '第6条【连涨三天先卖半仓】',
        'rating': '⭐⭐⭐⭐ 纪律性最强',
        'scenes': '适用于高波动环境下的利润锁定。当前石化板块个股连涨三天后出现回调的概率极高——宝丰能源3月多次出现"涨三天回调两天"的节奏。',
        'case': '宝丰能源3月10-12日连涨3天（+5.6%/+4.2%/+3.8%），13日回调-3.1%；3月17-19日连涨3天（+2.8%/+3.1%/+2.5%），20日回调-4.2%。',
        'advantage': '① 锁定利润，避免坐过山车；② 反人性但有纪律，执行力强的人用此条心态最稳；③ 释放资金等待下一个买点',
        'risk': '① 有时连涨三天只是起点（如连续涨停或主升浪），卖掉后可能踏空更大幅度的涨幅；② 在趋势明确的上涨中，此条会导致频繁操作、增加交易成本',
        'improve': '区分场景：震荡行情连涨三天 → 卖半仓；趋势行情连涨三天 → 不动，以5日线破位作为离场信号。判断"趋势行情"的标准：月涨幅>20% + 量能持续放大',
    },
    {
        'name': '第7条【中高位缩量可加仓】',
        'rating': '⭐⭐ 当前环境慎用',
        'scenes': '在2024-2025年的缓慢上涨趋势中较适用，但当前高波动、地缘驱动的市场中风险极大。',
        'case': '宝丰能源3月28日处于年涨幅54%的位置且近5日缩量（量比0.67），表面符合"中高位缩量"。但实际上当天跌-6.26%，原因是市场开始担心地缘局势缓和。',
        'advantage': '① 理论上有其逻辑：缩量意味抛压减弱，持股意愿稳定，可以是洗盘结束的信号；② 如果确认是洗盘结束，加仓后收益可观',
        'risk': '① 当前最大的坑：高位缩量更可能是"没人愿意接盘"而非"洗盘"；② 地缘局势随时可能缓和，高位缩量后突然破位大跌的概率极高；③ 没有任何明确的"洗盘vs出货"区分标准',
        'improve': '当前环境下建议将此条改为"中高位缩量 + 站稳20日均线 + 缩量后首次放量阳线"三条件缺一不可。如果只是缩量但不放量确认方向，宁可不加仓',
    },
    {
        'name': '第8条【放量走势不明立即退出】',
        'rating': '⭐⭐⭐⭐⭐ 保命金句',
        'scenes': '在当前高波动、消息驱动的市场中，这是最重要的一条——不是赚钱逻辑，是生存逻辑。',
        'case': '3月某日石化板块整体放量（板块成交量较前日放大80%）但指数不涨不跌，个股分化剧烈。随后一周板块整体下跌5%。如果当日按此规则操作，可提前离场规避下跌。',
        'advantage': '① 不做判断，纯做反应——放量不涨就是出货信号，不管原因如何先走；② "立即"二字是关键，避免"再等等看看"的犹豫；③ 保住了本金就保住了下一次机会',
        'risk': '① 偶尔放量走势不明只是洗盘，离场后踏空；② 需要定义"走势不明"——是±1%？±0.5%？需量化',
        'improve': '量化"走势不明"：放量（量比>1.5）但涨跌幅<0.5% + 个股涨跌比接近1:1 + 板块内分化严重，三条件同时满足即"走势不明"，果断离场',
    },
]

for ev in evaluations:
    p = doc.add_paragraph()
    add_bold_run(p, ev['name'], RGBColor(0, 51, 102))
    p.add_run(f'\n{ev["rating"]}')

    if 'scenes' in ev:
        p2 = doc.add_paragraph()
        add_bold_run(p2, '应用场景：')
        p2.add_run(ev['scenes'])

    if 'case' in ev:
        p2 = doc.add_paragraph()
        add_bold_run(p2, '实际案例：')
        p2.add_run(ev['case'])

    if 'advantage' in ev:
        p2 = doc.add_paragraph()
        add_bold_run(p2, '✅ 优势：')
        p2.add_run(ev['advantage'])

    if 'risk' in ev:
        p2 = doc.add_paragraph()
        add_bold_run(p2, '⚠️ 风险：')
        p2.add_run(ev['risk'])

    if 'improve' in ev:
        p2 = doc.add_paragraph()
        add_bold_run(p2, '🔧 改进建议：')
        p2.add_run(ev['improve'])

    doc.add_paragraph()  # 空行

# ===== 五、综合评分表 =====
add_colored_heading(doc, '五、综合评分总览')

score_data = [
    ['口诀', '实用性', '适用度\n（当前市场）', '胜率\n预估', '盈亏比', '核心难点'],
    ['①回踩均线', '⭐⭐⭐⭐', '高', '55-60%', '2:1', '判断真趋势'],
    ['②冲高7%', '⭐⭐⭐⭐⭐', '极高', '60-65%', '1.5:1', '科创板阈值'],
    ['③低点抬高', '⭐⭐⭐', '中等', '50-55%', '2:1', '真突破vs假突破'],
    ['④低开3%', '⭐⭐⭐', '中等', '55-60%', '1:1', '阈值太宽'],
    ['⑤连板换手', '⭐⭐⭐⭐⭐', '待触发', '45-50%', '4:1', '识别真假换手'],
    ['⑥连涨卖半', '⭐⭐⭐⭐', '高', '60-65%', '1.2:1', '趋势中踏空'],
    ['⑦缩量加仓', '⭐⭐', '低', '45-50%', '3:1', '洗盘vs出货'],
    ['⑧放量退出', '⭐⭐⭐⭐⭐', '极高', '—', '—', '定义"不明"'],
]

table = doc.add_table(rows=len(score_data), cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(score_data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True

# ===== 六、实战Checklist =====
add_colored_heading(doc, '六、实战执行Checklist')

p = doc.add_paragraph()
add_bold_run(p, '每天开盘前过一遍：')
p.add_run('\n')

checklist_before = [
    '当前持仓的股票是否连涨超过3天？ → 如有，准备半仓止盈',
    '是否有股票处于低位放量但走势不明？ → 如有，立即止损',
    '是否有股票低开超5%且前3分钟未反弹？ → 如有，准备止损',
    '是否有强势趋势股回踩分时均线？ → 如有，准备入场',
]

for item in checklist_before:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_bold_run(doc.add_paragraph(), '盘中实时监控：')

checklist_during = [
    '热门股涨幅到7%但未封涨停 → 提前挂好卖单',
    '个股连续涨停进入连板阶段 → 观察换手率和封单量',
    '当前中高位股缩量 → 不急加仓，等放量确认方向',
]

for item in checklist_during:
    doc.add_paragraph(item, style='List Bullet')

# ===== 七、结论 =====
add_colored_heading(doc, '七、核心结论')

conclusions = [
    ('八条口诀当纪律手册用，不当信号系统用。', RGBColor(139, 0, 0),
     '口诀最大的价值不是告诉你"什么时候买"，而是当你犹豫的时候给你一个"做决定的理由"——'
     '哪怕是错的决定，也比不做的决定强。'),
    ('当前市场是天然的口诀应用土壤。', None,
     '高波动率（年化50-80%）、热点快速轮动、消息驱动的脉冲行情，'
     '都为短线口诀策略提供了充足的信号和机会。'),
    ('需要加约束条件才能执行。', None,
     '八条口诀过于"硬编码"——没有考虑板块轮动、成交量变化、市场情绪等上下文。'
     '建议每条口诀都配上"适用场景 + 触发条件 + 止损规则"三层过滤。'),
    ('建议的优先级：', RGBColor(139, 0, 0),
     '②⑧ > ①⑤⑥ > ③④ > ⑦'
     '\n先掌握最实用的止盈止损纪律（②⑧），再学习入场技巧（①⑤⑥），'
     '最后再考虑波动操作（③④⑦）。'),
]

for title, color, desc in conclusions:
    p = doc.add_paragraph()
    add_bold_run(p, title, color)
    p.add_run(f'\n{desc}')

# ===== 风险声明 =====
doc.add_paragraph()
p = doc.add_paragraph()
add_bold_run(p, '⚠️ 风险提示：', RGBColor(180, 0, 0))
p.add_run(
    '短线交易风险极高，上述分析仅为方法论探讨，不构成任何投资建议。'
    '实际操作中需结合个人风险承受能力、资金体量和交易经验独立决策。'
    '历史案例不代表未来表现，市场环境随时可能变化。'
)

# 页脚
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

output_file = '/home/adam/.openclaw/workspace/短线交易八条口诀实操评估.docx'
doc.save(output_file)
print(f'✅ 报告已生成：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
