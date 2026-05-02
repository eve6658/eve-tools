#!/usr/bin/env python3
"""生成 A股政策市选股方法论分析报告 .docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(doc, text, color=RGBColor(139, 0, 0)):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = color

def bold(p, text, color=None):
    run = p.add_run(text)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    return run

def add_table(doc, data, style_name='Light Grid Accent 1'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = style_name
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
    return table

# ===== 封面 =====
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A股政策市选股方法论')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于养猪/饲料行业的半年行情实证检验')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run(f'报告日期：2026年3月29日')
run.font.size = Pt(12)

doc.add_page_break()

# ===== 一、核心论点 =====
add_heading(doc, '一、核心论点')

p = doc.add_paragraph()
bold(p, '一句话概括：')
p.add_run(
    'A股是政策市。对于涉及民生、三农的行业，选股逻辑的前提条件是"保供应、稳价格"，'
    '不能套用国际地缘政治的涨价逻辑。股票预期必须建立在这个先决条件之上。'
)

p = doc.add_paragraph()
bold(p, '反面案例：', RGBColor(180, 0, 0))
p.add_run(
    '霍尔木兹海峡封锁 → 国际化肥涨价 → 套用"涨价受益"逻辑买化肥股 → 云天化跌23%、兴发集团跌23%'
)

# ===== 二、行业分类框架 =====
add_heading(doc, '二、A股行业分类框架')

class_data = [
    ['行业类型', '特征', '政府干预程度', '选股逻辑', '代表行业'],
    ['政策稳价型', '涉及民生，价格是政治任务', '强干预（限价、储备投放、出口禁令）', '不做涨价预期，选成本领先+产能周期底部', '化肥、猪肉、粮食、电力、燃气'],
    ['市场定价型', '供需决定价格，企业享受弹性', '弱干预（基本不管）', '买涨价弹性+供给约束', '化工品、有色金属、半导体材料'],
    ['政策鼓励型', '政府推动产业升级', '正向干预（补贴、税收优惠）', '买政策红利+技术壁垒', '新能源、半导体、AI'],
]

add_table(doc, class_data)

p = doc.add_paragraph()
p.space_before = Pt(12)
bold(p, '关键判断步骤：')
p.add_run('\n① 先问：这个行业是不是民生/三农？')
p.add_run('\n② 再问：政策允许涨价吗？')
p.add_run('\n③ 如果答案是"不允许"，就不做涨价预期，只做成本分析和周期底部判断。')

# ===== 三、霍尔木兹危机下的案例对比 =====
add_heading(doc, '三、霍尔木兹危机下的行业对比')

p = doc.add_paragraph('2026年2月28日霍尔木兹海峡封锁后，不同行业类型的股价表现截然不同：')

compare_data = [
    ['行业', '类型', '3月均涨跌', '典型公司', '涨跌', '逻辑'],
    ['煤化工', '市场定价型', '↑ +11~15%', '宝丰能源', '+14.41%', '原料自给+产品涨价=利润扩张'],
    ['半导体特气', '政策鼓励型', '↑ +30~44%', '华特气体', '+44.24%', '国产替代+半导体景气=政策红利'],
    ['纯化肥', '政策稳价型', '↓ -18~25%', '云天化', '-23.26%', '限价+禁出口+成本涨=三重挤压'],
    ['养猪', '政策稳价型', '↓ -2~10%', '牧原股份', '-2.80%', '猪价被削峰填谷+成本涨=利润薄'],
    ['饲料', '政策稳价型', '↓ -3~7%', '海大集团', '-3.72%', '原料涨+成品不涨=利润被挤压'],
]

add_table(doc, compare_data)

# ===== 四、养猪 & 饲料行业验证 =====
add_heading(doc, '四、养猪 & 饲料行业半年行情验证')

doc.add_heading('4.1 养猪板块 — 过去6个月表现', level=2)

pig_data = [
    ['代码', '公司', '最新价', '6个月', '3个月', '1个月', '最大回撤'],
    ['002157.SZ', '正邦科技', '3.09', '+3.34%', '-9.12%', '-2.83%', '-18.50%'],
    ['300761.SZ', '立华股份', '20.47', '-9.42%', '+1.29%', '+0.39%', '-17.38%'],
    ['300498.SZ', '温氏股份', '16.91', '-10.10%', '+1.14%', '+6.76%', '-19.40%'],
    ['002714.SZ', '牧原股份', '45.20', '-12.74%', '-6.03%', '-2.80%', '-18.99%'],
    ['002124.SZ', '天邦食品', '2.40', '-18.64%', '-13.98%', '-4.00%', '-22.18%'],
    ['603363.SH', '傲农生物', '3.94', '-20.56%', '-14.35%', '-5.29%', '-33.62%'],
]

add_table(doc, pig_data)

p = doc.add_paragraph()
bold(p, '6个月均值：-11.35%')
p.add_run('（6只中5只下跌，仅正邦科技微涨3.34%）')

doc.add_heading('4.2 饲料板块 — 过去6个月表现', level=2)

feed_data = [
    ['代码', '公司', '最新价', '6个月', '3个月', '1个月', '最大回撤'],
    ['603299.SH', '苏盐井神', '11.82', '+11.40%', '+14.98%', '-1.66%', '-19.48%'],
    ['002840.SZ', '华统股份', '11.36', '+3.74%', '+20.21%', '+14.63%', '-21.10%'],
    ['002100.SZ', '天康生物', '7.30', '-3.18%', '+1.11%', '+2.24%', '-13.31%'],
    ['002688.SZ', '金河生物', '5.90', '-10.33%', '-2.64%', '-5.30%', '-20.82%'],
    ['603609.SH', '禾丰股份', '6.84', '-15.66%', '-3.66%', '-7.19%', '-18.94%'],
    ['002311.SZ', '海大集团', '52.06', '-17.87%', '-4.04%', '-3.72%', '-21.26%'],
]

add_table(doc, feed_data)

p = doc.add_paragraph()
bold(p, '6个月均值：-5.32%')
p.add_run('（6只中4只下跌）')

doc.add_heading('4.3 养猪+饲料一体化', level=2)

combo_data = [
    ['代码', '公司', '最新价', '6个月', '3个月', '1个月', '最大回撤'],
    ['002548.SZ', '金新农', '5.55', '+10.34%', '-8.11%', '-20.14%', '-28.69%'],
    ['002385.SZ', '大北农', '3.96', '-2.94%', '-1.25%', '-1.98%', '-16.56%'],
    ['000876.SZ', '新希望', '8.24', '-15.05%', '-9.85%', '-4.85%', '-19.20%'],
    ['002567.SZ', '唐人神', '4.05', '-15.80%', '-8.37%', '-6.68%', '-18.89%'],
]

add_table(doc, combo_data)

p = doc.add_paragraph()
bold(p, '6个月均值：-5.86%')
p.add_run('（4只中3只下跌）')

# ===== 五、验证结论 =====
add_heading(doc, '五、验证结论')

summary = [
    ['行业分类', '6个月均值', '下跌比例', '验证结果'],
    ['养猪', '-11.35%', '83%', '❌ 完全符合"政策稳价型"判断'],
    ['饲料', '-5.32%', '67%', '❌ 部分符合（2只逆势涨，因非纯饲料业务）'],
    ['养猪+饲料', '-5.86%', '75%', '❌ 符合判断'],
    ['纯化肥', '-18~25%', '100%', '❌ 完全符合（限价+禁出口）'],
    ['煤化工', '+11~15%', '0%下跌', '✅ 符合"市场定价型"判断'],
]

add_table(doc, summary)

p = doc.add_paragraph()
p.space_before = Pt(12)
bold(p, '关键发现：', RGBColor(139, 0, 0))
p.add_run(
    '\n① 政策稳价型行业（养猪、饲料、化肥）过去6个月整体亏损，验证框架正确'
    '\n② 养猪比饲料更差（-11% vs -5%），因为养猪受猪价直接压制，饲料只是间接传导'
    '\n③ 最大回撤普遍18-34%，持有体验极差'
    '\n④ 例外存在：苏盐井神(+11.40%)和华统股份(+3.74%)逆势上涨，原因是它们有盐业/屠宰等非饲料主业'
)

# ===== 六、方法论总结 =====
add_heading(doc, '六、政策市选股方法论')

p = doc.add_paragraph()
bold(p, 'Step 1：判断行业类型')
p.add_run(
    '\n→ 问自己：这个行业涉及民生/三农吗？'
    '\n→ 如果是，属于"政策稳价型"'
    '\n→ 如果不是，属于"市场定价型"或"政策鼓励型"'
)

p = doc.add_paragraph()
bold(p, 'Step 2：问"政策允许涨价吗"')
p.add_run(
    '\n→ 不允许 → 不做涨价预期，只做成本分析'
    '\n→ 允许 → 做涨价弹性分析'
)

p = doc.add_paragraph()
bold(p, 'Step 3：选股方向')
p.add_run(
    '\n→ 政策稳价型：选成本领先 + 产能周期底部 + 高股息防御'
    '\n→ 市场定价型：选涨价弹性 + 供给约束 + 原料自给'
    '\n→ 政策鼓励型：选技术壁垒 + 政策补贴 + 国产替代'
)

p = doc.add_paragraph()
bold(p, 'Step 4：风险预警信号')
p.add_run(
    '\n→ 政府出台限价政策 → 政策稳价型行业利空'
    '\n→ 国储大规模投放 → 供给增加=价格承压'
    '\n→ 出口禁令 → 收入渠道被砍=直接利空'
    '\n→ 国际涨价新闻 → 先判断：政策允许传导吗？不允许就只是噪音'
)

# ===== 七、反面教材 =====
add_heading(doc, '七、反面教材——霍尔木兹危机中的散户陷阱')

traps = [
    ('陷阱一："国际涨价=国内涨"',
     '散户看到"国际化肥涨80%"就冲进去买化肥股，不知道政策已经限价+禁出口。结果云天化跌23%。'),
    ('陷阱二："涨价利好=公司受益"',
     '猪价涨→政府投放储备→猪价回落。养猪股的利润弹性被政策"削峰填谷"。'),
    ('陷阱三："板块龙头=最佳标的"',
     '海大集团是饲料龙头，但6个月跌17.87%。苏盐井神是边缘公司，但涨11.40%。'
     '原因：苏盐井神有盐业业务不受饲料政策影响。板块内选股要看业务结构，不能只看龙头。'),
]

for title, desc in traps:
    p = doc.add_paragraph()
    bold(p, f'【{title}】')
    p.add_run(f'\n{desc}')

# ===== 结论 =====
add_heading(doc, '八、核心结论')

conclusions = [
    'A股政策市选股框架经半年行情实证检验，结论可靠。',
    '民生/三农行业（养猪、饲料、化肥）整体跑输市场，持有体验极差。',
    '唯一的"例外受益者"是煤化工——它同时享受"化工品涨价"（市场定价）+ "原料自给"（不受进口影响）。',
    '选股的第一步不是选公司，而是先判断行业属于哪一类。选错类型，再好的公司也难赚钱。',
]

for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    bold(p, f'{i}. ', RGBColor(139, 0, 0))
    p.add_run(c)

# 风险声明
doc.add_paragraph()
p = doc.add_paragraph()
bold(p, '⚠️ 风险提示：', RGBColor(180, 0, 0))
p.add_run('本报告为方法论探讨，基于历史数据分析，不构成投资建议。投资者应根据自身情况独立决策。')

# 页脚
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

output_file = '/home/adam/.openclaw/workspace/A股政策市选股方法论分析报告.docx'
doc.save(output_file)
print(f'✅ 报告已生成：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
