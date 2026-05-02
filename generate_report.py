#!/usr/bin/env python3
"""生成 A股石油化工板块 2026年3月行情分析报告 .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd
import os
from datetime import datetime

# 读取数据
csv_file = '/home/adam/.openclaw/workspace/petrochemical_march_2026.csv'
df = pd.read_csv(csv_file)

# 公司名称映射
name_map = {
    '600989.SH': '宝丰能源', '601898.SH': '中煤能源', '002490.SZ': '山东墨龙',
    '601857.SH': '中国石油', '002267.SZ': '陕天然气', '002221.SZ': '华油能源',
    '600546.SH': '山煤国际', '601001.SH': '大同煤业', '000698.SZ': '沈阳化工',
    '600232.SH': '金鹰股份', '002493.SZ': '荣盛石化', '002207.SZ': '准油股份',
    '002278.SZ': '神开股份', '600346.SH': '恒力石化', '002554.SZ': '惠博普',
    '600688.SH': '上海石化', '601808.SH': '中海油服', '600583.SH': '海油工程',
    '600871.SH': '石化油服', '600968.SH': '海油发展', '603619.SH': '中曼石油',
    '600028.SH': '中国石化', '000059.SZ': '华锦股份', '002629.SZ': '仁智股份',
    '000096.SZ': '广聚能源', '002476.SZ': '宝莫股份', '601001.SH': '大同煤业',
    '002778.SZ': '中峻能源',
}

def get_name(code):
    return name_map.get(code, code)

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 封面标题 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run('A股石油化工板块')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年3月行情分析报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run(f'报告日期：2026年3月28日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== 一、摘要 =====
doc.add_heading('一、报告摘要', level=1)

p = doc.add_paragraph()
p.add_run('核心结论：').bold = True
p.add_run('2026年3月，A股石油化工板块整体表现疲软，29只样本股中仅4只实现正收益，板块平均涨幅为-8.81%，中位数涨幅-11.04%。'
           '宝丰能源（600989.SH）以+14.41%的涨幅领跑板块，中煤能源（601898.SH）紧随其后，涨幅达+11.75%。'
           '中国石油（601857.SH）作为权重股逆势上涨+2.29%，发挥了板块"稳定器"作用。')

doc.add_heading('关键数据速览', level=2)

# 汇总数据表
summary_data = [
    ['指标', '数值'],
    ['分析股票数', '29 只'],
    ['3月上涨', f"{(df['月涨跌幅(%)'] > 0).sum()} 只（{(df['月涨跌幅(%)'] > 0).sum()/len(df)*100:.1f}%）"],
    ['3月下跌', f"{(df['月涨跌幅(%)'] < 0).sum()} 只（{(df['月涨跌幅(%)'] < 0).sum()/len(df)*100:.1f}%）"],
    ['平均涨幅', f"{df['月涨跌幅(%)'].mean():.2f}%"],
    ['中位数涨幅', f"{df['月涨跌幅(%)'].median():.2f}%"],
    ['最大涨幅', f"{df['月涨跌幅(%)'].max():.2f}%"],
    ['最大跌幅', f"{df['月涨跌幅(%)'].min():.2f}%"],
]

table = doc.add_table(rows=len(summary_data), cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(summary_data):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(11)
                if i == 0:
                    run.font.bold = True

# ===== 二、涨幅TOP 10 =====
doc.add_heading('二、3月涨幅TOP 10', level=1)

p = doc.add_paragraph('以下为2026年3月A股石油化工板块涨幅前十名：')

# TOP10 表格
top10 = df.head(10)
top10_headers = ['排名', '代码', '公司名称', '月涨跌幅(%)', '月末收盘价', '振幅(%)', '上涨天数', '下跌天数', '日均成交额(亿)', '交易天数']

table = doc.add_table(rows=len(top10)+1, cols=len(top10_headers))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
for j, header in enumerate(top10_headers):
    cell = table.cell(0, j)
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# 数据行
for i, (_, row) in enumerate(top10.iterrows()):
    values = [
        str(i + 1),
        row['ts_code'],
        get_name(row['ts_code']),
        f"{row['月涨跌幅(%)']:+.2f}%",
        f"{row['月末收盘']:.2f}",
        f"{row['振幅(%)']:.2f}%",
        str(row['上涨天数']),
        str(row['下跌天数']),
        f"{row['日均成交额(亿)']:.2f}",
        str(row['交易天数']),
    ]
    for j, val in enumerate(values):
        cell = table.cell(i + 1, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                # 涨幅正数标红，负数标绿
                if j == 3 and row['月涨跌幅(%)'] > 0:
                    run.font.color.rgb = RGBColor(220, 20, 20)
                elif j == 3 and row['月涨跌幅(%)'] < 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)

# ===== 三、个股点评 =====
doc.add_heading('三、重点个股点评', level=1)

highlights = [
    ('宝丰能源（600989.SH）', '+14.41%',
     '3月表现最佳，涨幅达14.41%，振幅42.80%显示市场关注度极高。'
     '作为煤制烯烃龙头企业，受益于能源安全战略和煤化工景气度回升。'
     '全月上涨天数11天，多于下跌天数，走势稳健。'),
    ('中煤能源（601898.SH）', '+11.75%',
     '涨幅位居第二，振幅25.94%。作为大型煤炭央企，受益于煤炭板块整体回暖。'
     '全月上涨下跌各10天，走势相对均衡但趋势向上。'),
    ('中国石油（601857.SH）', '+2.29%',
     '作为板块绝对权重股，在板块整体下跌8.81%的背景下逆势上涨2.29%，'
     '堪称板块"定海神针"。总成交额1.20亿元，流动性充裕。'
     '上涨11天多于下跌9天，走势偏强。'),
    ('山东墨龙（002490.SZ）', '+4.24%',
     '涨幅+4.24%位列第三，但振幅高达51.86%，波动极大。'
     '上涨12天，下跌8天，市场分歧较大，建议关注后续催化因素。'),
]

for name, chg, comment in highlights:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}  |  3月涨跌幅：{chg}')
    run.font.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(comment)
    p.paragraph_format.space_after = Pt(12)

# ===== 四、板块整体分析 =====
doc.add_heading('四、板块整体分析', level=1)

doc.add_heading('4.1 涨跌分布', level=2)
p = doc.add_paragraph()
p.add_run('3月石油化工板块呈现明显的普跌格局：')
p.add_run(f'\n• 上涨股票仅 {int((df["月涨跌幅(%)"] > 0).sum())} 只，占比 {(df["月涨跌幅(%)"] > 0).sum()/len(df)*100:.1f}%')
p.add_run(f'\n• 下跌股票 {int((df["月涨跌幅(%)"] < 0).sum())} 只，占比 {(df["月涨跌幅(%)"] < 0).sum()/len(df)*100:.1f}%')
p.add_run(f'\n• 平均跌幅 {df["月涨跌幅(%)"].mean():.2f}%，中位数跌幅 {df["月涨跌幅(%)"].median():.2f}%')

doc.add_heading('4.2 跌幅前5', level=2)

bottom5 = df.tail(5).sort_values('月涨跌幅(%)', ascending=True)
table = doc.add_table(rows=len(bottom5)+1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, header in enumerate(['代码', '公司名称', '月涨跌幅(%)', '振幅(%)']):
    cell = table.cell(0, j)
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for i, (_, row) in enumerate(bottom5.iterrows()):
    vals = [row['ts_code'], get_name(row['ts_code']), f"{row['月涨跌幅(%)']:+.2f}%", f"{row['振幅(%)']:.2f}%"]
    for j, val in enumerate(vals):
        cell = table.cell(i + 1, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_heading('4.3 成交活跃度', level=2)
p = doc.add_paragraph()
total_amt = df['总成交额(亿)'].sum()
top_amt_code = df.loc[df['总成交额(亿)'].idxmax(), 'ts_code']
p.add_run(f'板块3月总成交额约 {total_amt:.2f} 亿元。'
          f'成交额最高的是 {get_name(top_amt_code)}（{top_amt_code}），'
          f'达 {df["总成交额(亿)"].max():.2f} 亿元。')

# ===== 五、风险提示 =====
doc.add_heading('五、风险提示', level=1)

risks = [
    '本报告基于历史数据分析，不构成投资建议。',
    '由于Tushare行业分类接口权限限制，成分股列表基于已知核心石油石化企业构建，可能未完全覆盖全部细分个股。',
    '石油化工行业受国际原油价格、地缘政治、政策调控等多重因素影响，投资需注意相关风险。',
    '数据截至2026年3月28日，后续行情可能存在变化。',
]
for risk in risks:
    doc.add_paragraph(risk, style='List Bullet')

# ===== 页脚信息 =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

# 保存
output_file = '/home/adam/.openclaw/workspace/石油化工板块_2026年3月行情分析报告.docx'
doc.save(output_file)
print(f'✅ 报告已生成：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
