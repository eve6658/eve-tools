#!/usr/bin/env python3
"""生成 霍尔木兹海峡危机对中国石油化工煤化工行业影响分析报告 .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime

doc = Document()

# 默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(100)
run = p.add_run('霍尔木兹海峡危机')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('对中国石油化工/煤化工行业影响')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——暨2026年4月投资展望')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(60)
run = p.add_run('2026年3月29日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== 一、摘要 =====
doc.add_heading('一、报告摘要', level=1)
p = doc.add_paragraph(
    '2026年2月28日，美以联合打击伊朗，伊朗革命卫队随即宣布霍尔木兹海峡禁航。'
    '截至3月中旬，通过海峡的油轮从日均24艘骤降至约4艘，全球原油日供应量骤降约800万桶（占全球8-10%），'
    '布伦特原油一度逼近120美元/桶。国际能源署称全球石油市场面临"史上最严重供应梗阻"。\n\n'
    '本报告分析这一危机对中国石油化工及煤化工行业的影响，并对2026年4月相关板块投资前景做出展望。'
)

p = doc.add_paragraph()
p.add_run('核心结论：').bold = True
p.add_run(
    '中国煤化工企业是此轮危机中罕见的"确定性受益"标的。煤化工以国内煤炭为原料，完全不受海运断供影响，'
    '同时享受产品价格飙升带来的利润弹性。4月板块走势高度依赖霍尔木兹海峡何时恢复通航。'
)

# ===== 二、局势背景 =====
doc.add_heading('二、当前局势背景', level=1)

doc.add_heading('2.1 事件时间线', level=2)
timeline = [
    ('2026年2月28日', '美以联合军事打击伊朗'),
    ('2028日当晚', '伊朗革命卫队宣布霍尔木兹海峡全面禁航'),
    ('3月3日', '革命卫队称十余艘油轮遭炮弹击毁，海峡进入"战争状态"'),
    ('3月4日', '布伦特原油突破84美元/桶，WTI突破76美元/桶'),
    ('3月11日', 'IEA成员国一致同意释放4亿桶战略石油储备'),
    ('3月12日', 'IEA月度报告：全球石油市场面临"史上最严重供应梗阻"'),
    ('3月27日', 'IRGC宣布海峡对美、以色列及其盟友港口船只全面关闭'),
]

table = doc.add_table(rows=len(timeline)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.cell(0, 0).text = '日期'
table.cell(0, 1).text = '事件'
for paragraph in table.cell(0, 0).paragraphs:
    for run in paragraph.runs:
        run.font.bold = True
for paragraph in table.cell(0, 1).paragraphs:
    for run in paragraph.runs:
        run.font.bold = True
for i, (date, event) in enumerate(timeline):
    table.cell(i+1, 0).text = date
    table.cell(i+1, 1).text = event

doc.add_heading('2.2 关键影响数据', level=2)

impact_data = [
    ['指标', '危机前', '危机后', '变化'],
    ['日均通过海峡油轮数', '24艘', '约4艘', '↓83%'],
    ['全球原油日供应量', '约1.02亿桶', '约9400万桶', '↓800万桶（-8%）'],
    ['布伦特原油价格', '~75美元/桶', '~120美元/桶', '↑60%'],
    ['中东炼油停运产能', '—', '300万桶/日以上', '显著'],
    ['海湾国家日减产', '—', '1000万桶+', '—'],
    ['IEA释放战略储备', '—', '4亿桶', '权宜之计'],
]

table = doc.add_table(rows=len(impact_data), cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(impact_data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

# ===== 三、影响逻辑链 =====
doc.add_heading('三、影响传导逻辑', level=1)

p = doc.add_paragraph()
p.add_run('此次危机对中国石化行业的影响可通过以下逻辑链传导：').bold = False

chain = [
    '霍尔木兹海峡封锁 → 波斯湾原油/LPG/石化原料断供',
    '→ 日本、韩国、越南、泰国等高度依赖中东原油的国家石化装置因原料不足被迫降负荷甚至停车',
    '→ 亚洲石化产品供给出现巨大缺口，产品价格飙升',
    '→ 中国煤化工企业（原料为国内煤炭，完全自给）成本不变但产品涨价 → 利润空间暴增',
    '→ 中国拥有原料库存的炼化企业及国内油气开采企业同步受益',
]

for i, step in enumerate(chain, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'第{i}步：{step}')
    if i == len(chain):
        run.font.bold = True
        run.font.color.rgb = RGBColor(139, 0, 0)

# ===== 四、受益分析 =====
doc.add_heading('四、受益公司分层分析', level=1)

doc.add_heading('4.1 第一梯队：煤化工（最大受益者）', level=2)

p = doc.add_paragraph(
    '煤化工以国内煤炭为原料，完全不受国际原油海运断供影响。'
    '当全球油制化工品成本随油价暴涨时，煤制化工品的成本几乎不变，利润空间大幅扩张。'
    '同时，日韩东南亚石化产能因缺油降负荷，订单将进一步转向中国煤化工企业。'
)

coal_chem = [
    ['公司', '代码', '核心优势', '主要产品'],
    ['宝丰能源', '600989.SH', '煤制烯烃龙头，产能领先，成本控制优秀', '聚乙烯、聚丙烯、焦炭'],
    ['中煤能源', '601898.SH', '煤炭+煤化工一体化，资源自给', '聚烯烃、甲醇、尿素'],
    ['华鲁恒升', '600426.SH', '煤化工多联产龙头，产品多元化', 'DMF、醋酸、己二酸'],
    ['卫星化学', '002648.SZ', 'C2/C3产业链完整', '丙烯酸、乙烯、聚乙烯'],
    ['宝泰隆', '601011.SH', '煤焦化深加工', '甲醇、煤焦油、针状焦'],
]

table = doc.add_table(rows=len(coal_chem), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(coal_chem):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

doc.add_heading('4.2 第二梯队：国内油气开采', level=2)

p = doc.add_paragraph(
    '油价高企直接利好国内油气生产企业，其产品售价跟随国际油价上涨，'
    '而开采成本相对固定，利润弹性显著。同时，能源安全战略下国内油气资本开支有望增加。'
)

oil_gas = [
    ['公司', '代码', '核心优势'],
    ['中国石油', '601857.SH', '国内最大油气生产商，油价直接受益，产量稳定增长'],
    ['中海油服', '601808.SH', '海上油气服务龙头，上游资本开支增加带动业绩'],
    ['海油工程', '600583.SH', '海上油气工程建设，受益于国内油气开发加速'],
    ['中国海油', '600938.SH', '纯上游油气生产商，油价弹性最大'],
]

table = doc.add_table(rows=len(oil_gas), cols=3)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(oil_gas):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

doc.add_heading('4.3 第三梯队：炼化一体化（受益但有压制）', level=2)

p = doc.add_paragraph(
    '大型炼化一体化企业受益于化工品涨价，但同时面临原料成本上升的压力。'
    '若企业前期有低价原油库存，则短期受益明显；若库存耗尽后原料需高价补充，利润将被压缩。'
)

refining = [
    ['公司', '代码', '特征'],
    ['荣盛石化', '002493.SZ', '浙石化4000万吨/年大炼化，有部分原料库存优势'],
    ['恒力石化', '600346.SH', '2000万吨/年炼化一体化，产业链完整'],
    ['上海石化', '600688.SH', '老牌炼化企业，产能规模适中'],
]

table = doc.add_table(rows=len(refining), cols=3)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(refining):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

# ===== 五、成本对比 =====
doc.add_heading('五、煤制vs油制化工品成本对比', level=1)

p = doc.add_paragraph('以下对比清晰展示煤化工在当前局势下的成本优势：')

cost_data = [
    ['指标', '油制路线（危机前）', '油制路线（危机后）', '煤制路线（当前）'],
    ['原料成本', '~75美元/桶原油', '~120美元/桶原油', '国产煤炭（基本稳定）'],
    ['聚乙烯成本', '约6000-6500元/吨', '约8500-9500元/吨', '约5000-6000元/吨'],
    ['聚丙烯成本', '约5800-6300元/吨', '约8000-9000元/吨', '约4800-5500元/吨'],
    ['甲醇成本', '约2500-3000元/吨', '约4000-5000元/吨', '约1800-2200元/吨'],
    ['利润空间', '正常', '被压缩（原料涨）', '大幅扩张（产品涨、成本稳）'],
]

table = doc.add_table(rows=len(cost_data), cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(cost_data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True

# ===== 六、3月市场表现回顾 =====
doc.add_heading('六、3月市场表现回顾', level=1)

p = doc.add_paragraph(
    '值得注意的是，3月份石油化工板块整体下跌8.81%的背景下，'
    '煤化工和油气开采龙头已逆势上涨，显示"聪明资金"已提前布局：'
)

march_data = [
    ['公司', '代码', '3月涨跌幅', '说明'],
    ['宝丰能源', '600989.SH', '+14.41%', '板块涨幅第一，振幅42.80%，资金高度关注'],
    ['中煤能源', '601898.SH', '+11.75%', '煤炭+煤化工双逻辑驱动'],
    ['中国石油', '601857.SH', '+2.29%', '大盘权重股逆势飘红，"定海神针"'],
    ['山东墨龙', '002490.SZ', '+4.24%', '油服设备概念，振幅高达51.86%'],
]

table = doc.add_table(rows=len(march_data), cols=4)
table.style = 'Light Grid Accent 1'
for i, row_data in enumerate(march_data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

# ===== 七、4月展望 =====
doc.add_heading('七、2026年4月投资展望', level=1)

doc.add_heading('7.1 情景分析', level=2)

scenarios = [
    ['情景', '概率评估', '油价区间', '煤化工板块预期'],
    ['冲突持续，海峡未恢复', '中高', '100-140美元', '强烈看多，估值重估，涨幅加速'],
    ['部分恢复，有限通航', '中等', '85-110美元', '中性偏多，受益程度减弱'],
    ['快速缓和，全面恢复通航', '较低', '70-85美元', '利空，板块回调风险大'],
]

table = doc.add_table(rows=len(scenarios), cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(scenarios):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
                if i == 1 and j == 3:
                    run.font.color.rgb = RGBColor(139, 0, 0)

doc.add_heading('7.2 分阶段判断', level=2)

p = doc.add_paragraph()
run = p.add_run('短期（4月上旬）：强烈看多 🔴')
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(139, 0, 0)
p = doc.add_paragraph(
    '• 地缘冲突持续发酵，市场情绪高涨\n'
    '• 3月煤化工板块已逆势上涨，显示资金持续流入\n'
    '• 若4月局势未缓和，涨幅有望加速'
)

p = doc.add_paragraph()
run = p.add_run('中期（4月中下旬）：取决于冲突走向 ⚡')
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(200, 150, 0)
p = doc.add_paragraph(
    '• 持续封锁 → 煤化工对标全球油制化工品利润差，估值体系重估\n'
    '• 部分恢复 → 油价回落但仍高位，受益逻辑减弱\n'
    '• 快速缓和 → 前期追高资金撤离，板块回调风险大'
)

doc.add_heading('7.3 值得关注的催化剂', level=2)
catalysts = [
    '霍尔木兹海峡局势进一步升级（禁航范围扩大、冲突升级）',
    'IEA战略储备释放不足以弥补供应缺口',
    '亚洲日韩等国石化装置大面积停车消息',
    '中国煤化工龙头一季报超预期（预计4月披露）',
    '国内政策层面出台能源安全保障措施',
    '化工品现货价格持续攀升',
]
for c in catalysts:
    doc.add_paragraph(c, style='List Bullet')

# ===== 八、风险提示 =====
doc.add_heading('八、风险提示', level=1)

risks = [
    '地缘政治风险：冲突缓和导致油价回落，前期追高资金可能被套。中国若释放战略储备，也将平抑油价涨幅。',
    '需求侧风险：若全球经济因高油价陷入衰退，化工品终端需求也会随之下滑。',
    '成本传导风险：煤炭价格可能因整体能源紧张而上涨，压缩煤化工利润空间。',
    '估值拥挤风险：板块交易拥挤度快速上升，短期涨幅过大存在回调压力。',
    '政策风险：政府可能出台价格管制或出口限制措施，影响企业利润。',
    '本报告基于公开信息分析，不构成投资建议。投资有风险，决策需谨慎。',
]
for risk in risks:
    doc.add_paragraph(risk, style='List Bullet')

# ===== 九、结论 =====
doc.add_heading('九、总结', level=1)

p = doc.add_paragraph()
p.add_run('核心结论：').bold = True
p.add_run(
    '霍尔木兹海峡危机使中国煤化工企业成为全球石化行业中罕见的"确定性受益"标的。'
    '在亚洲各国石化产能因原料断供而被迫降负荷之际，中国煤化工以国内煤炭为原料，'
    '完全不受海运危机影响，同时享受产品价格飙升带来的巨大利润弹性。\n\n'
    '4月投资策略建议关注两条主线：\n'
)
p.add_run('主线一（首选）：煤化工龙头 — ').bold = True
p.add_run('宝丰能源、华鲁恒升、中煤能源，成本优势确定，利润弹性最大。\n')
p.add_run('主线二（辅助）：国内油气开采 — ').bold = True
p.add_run('中国石油、中国海油，油价直接受益，防御性较好。\n\n')
p.add_run('关键变量：').bold = True
p.add_run('霍尔木兹海峡何时恢复通航，将直接决定行情持续时间。')

# 页脚
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

# 保存
output_file = '/home/adam/.openclaw/workspace/霍尔木兹海峡危机_中国石化煤化工影响分析报告.docx'
doc.save(output_file)
print(f'✅ 报告已生成：{output_file}')
print(f'📄 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB')
